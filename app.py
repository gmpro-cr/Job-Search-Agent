"""
app.py - Flask web UI for Job Search Agent.
Provides a browser-based interface for managing preferences, running the scraper,
viewing jobs, and browsing digests.
"""

import os
import sys
import uuid
import json
import logging
import threading
from datetime import datetime, timedelta

# Load .env before anything else reads os.environ
try:
    from dotenv import load_dotenv
    load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))
except ImportError:
    pass

from flask import (
    Flask, render_template, request, redirect, url_for,
    flash, jsonify, send_from_directory,
)

# Ensure project root is on the path so we can import sibling modules
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)

from main import load_config, load_preferences, save_preferences, DEFAULT_PREFS, apply_env_overrides, _CREDENTIAL_KEYS
from database import (
    init_db, get_connection, get_comprehensive_stats, get_portal_quality_stats,
    update_applied_status, insert_jobs_bulk, generate_job_id, mark_sent_in_digest,
    select_digest_jobs,
    get_unsent_jobs, update_job_contacts, get_distinct_locations,
    get_normalized_locations, normalize_location, _CITY_PATTERNS,
    get_application_pipeline_stats, get_best_matching_categories,
    get_application_activity, get_recommended_actions,
    hide_job, update_job_notes, dedup_jobs,
    _INTERNATIONAL_CANONICALS, _INTERNATIONAL_KEYWORDS,
    # Phase 1 multi-user helpers
    get_or_create_user, get_user_by_email,
    update_applied_status_user, update_job_notes_user, hide_job_user, mark_job_viewed,
    bulk_set_user_cv_scores,
    get_comprehensive_stats_user, get_dashboard_insights_user,
    get_application_pipeline_stats_user_with_legacy_fallback,
    user_state_join_sql,
    is_admin_user, promote_first_user_to_admin,
    get_user_cv_data, purge_stale_demo_users,
)
from scrapers import scrape_all_portals
from analyzer import analyze_jobs, generate_tailored_points, parse_nlp_query, parse_cv_text, cv_score, compute_gap_analysis, load_cv_data, save_cv_data, CV_DATA_PATH
from digest_generator import generate_digest, get_latest_digest, DIGEST_DIR
from email_notifier import send_job_email
from database import delete_old_jobs
from contact_scraper import enrich_jobs_with_contacts
from database import update_job_description
from telegram_notifier import send_telegram_alert, send_telegram_batch_summary
from telegram_bot import start_telegram_bot
from git_sync import sync_from_scrape

# ---------------------------------------------------------------------------
# App setup
# ---------------------------------------------------------------------------

_IS_VERCEL = bool(os.environ.get("VERCEL"))
# Dev login is OFF by default everywhere; opt in with ENABLE_DEV_LOGIN=1
# (intended for local dev only). It bypasses OAuth so we never want it
# accidentally enabled on a hosted deployment.
_ENABLE_DEV_LOGIN = os.environ.get("ENABLE_DEV_LOGIN") == "1" and not _IS_VERCEL

# Demo login lets first-time visitors explore the product flow WITHOUT Google
# sign-in. Unlike dev login it is allowed on hosted deployments — it's a
# temporary, public "try it" door. Each demo login creates a fresh throwaway
# user (demo-*@demo.local) so every session starts as a genuine first-timer.
# ON by default; set ENABLE_DEMO_LOGIN=0 to switch it off without a redeploy.
_ENABLE_DEMO_LOGIN = os.environ.get("ENABLE_DEMO_LOGIN", "1") != "0"
_DEMO_EMAIL_DOMAIN = "demo.local"

app = Flask(__name__)

_secret = os.environ.get("FLASK_SECRET")
if not _secret:
    if _IS_VERCEL:
        raise RuntimeError(
            "FLASK_SECRET env var is required in production. "
            "Generate with: python -c 'import secrets; print(secrets.token_urlsafe(48))'"
        )
    _secret = "dev-only-insecure-key-do-not-use-in-production"
app.secret_key = _secret

app.config["SESSION_COOKIE_SECURE"]   = _IS_VERCEL
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=7)
# 4 MB is more than enough for a CV (PDF/DOCX). Smaller cap reduces
# memory pressure under abuse.
app.config["MAX_CONTENT_LENGTH"] = 4 * 1024 * 1024

# Logging: stderr on Vercel (platform collects it), rotating file locally.
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s")
logger = logging.getLogger(__name__)
if not _IS_VERCEL:
    try:
        from logging.handlers import RotatingFileHandler
        _log_path = os.path.join(BASE_DIR, "app.log")
        _file_handler = RotatingFileHandler(
            _log_path, maxBytes=2 * 1024 * 1024, backupCount=3, encoding="utf-8"
        )
        _file_handler.setFormatter(logging.Formatter(
            "%(asctime)s [%(levelname)s] %(name)s: %(message)s"
        ))
        # Attach to the root so module loggers inherit it.
        logging.getLogger().addHandler(_file_handler)
    except Exception as _e:
        logger.warning("RotatingFileHandler setup failed: %s", _e)


@app.after_request
def _security_headers(response):
    """Conservative, defence-in-depth headers."""
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "DENY")
    response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    response.headers.setdefault("Permissions-Policy",
                                "geolocation=(), microphone=(), camera=()")
    if _IS_VERCEL:
        response.headers.setdefault(
            "Strict-Transport-Security",
            "max-age=63072000; includeSubDomains"
        )
    return response

# ---------------------------------------------------------------------------
# CSRF protection
# ---------------------------------------------------------------------------
from flask_wtf.csrf import CSRFProtect, generate_csrf
csrf = CSRFProtect(app)

@app.after_request
def set_csrf_cookie(response):
    response.headers['X-CSRF-Token'] = generate_csrf()
    return response

# ---------------------------------------------------------------------------
# Auth setup
# ---------------------------------------------------------------------------
from authlib.integrations.flask_client import OAuth
from flask import session

oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=os.environ.get("GOOGLE_CLIENT_ID"),
    client_secret=os.environ.get("GOOGLE_CLIENT_SECRET"),
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={'scope': 'openid email profile'}
)

_PUBLIC_ENDPOINTS = frozenset({
    'login', 'auth_google', 'auth_callback', 'static',
    'approve_outreach', 'skip_outreach', 'favicon',
    'index', 'privacy', 'terms',
} | ({'auth_dev_login'} if _ENABLE_DEV_LOGIN else set())
  | ({'auth_demo_login'} if _ENABLE_DEMO_LOGIN else set()))

@app.before_request
def require_login():
    if request.endpoint in _PUBLIC_ENDPOINTS:
        return
    if session.get('user') is None:
        if request.is_json or (request.path or '').startswith('/api/'):
            from flask import abort
            abort(401)
        return redirect(url_for('login', next=request.path))


def current_user_id():
    """
    Resolve the DB user id for the current session, creating the user row on
    first access. Returns None for unauthenticated routes. Callers that require
    a user should use require_user_id() which aborts with 401.
    """
    user = dict(session).get('user') or {}
    uid = user.get('id')
    if uid:
        try:
            return int(uid)
        except (TypeError, ValueError):
            pass
    email = (user.get('email') or '').strip()
    if not email:
        return None
    try:
        uid = get_or_create_user(email, user.get('name'), user.get('picture'))
    except Exception as e:
        logger.warning("get_or_create_user failed for %s: %s", email, e)
        return None
    # Cache on the session so we don't hit the DB on every request
    session['user'] = dict(user, id=uid)
    return uid


def require_user_id():
    """current_user_id() that aborts the request with 401 when no user."""
    uid = current_user_id()
    if not uid:
        from flask import abort
        abort(401)
    return uid


def _is_demo_user():
    """True when the current session is a throwaway demo user (no Google)."""
    user = dict(session).get('user') or {}
    if user.get('demo'):
        return True
    return (user.get('email') or '').endswith(f"@{_DEMO_EMAIL_DOMAIN}")


def require_admin():
    """Like require_user_id() but also enforces is_admin=1. 403 if not."""
    uid = require_user_id()
    if not is_admin_user(uid):
        from flask import abort
        abort(403)
    return uid
# Initialize the database on startup.
#
# init_db() is idempotent (CREATE TABLE IF NOT EXISTS) but issues ~10 DDL
# round-trips. On Vercel that's wasted latency on every cold start, since
# the schema is already provisioned. Skip it when SKIP_INIT_DB is set;
# keep running it locally so dev environments bootstrap fresh tables.
if not os.environ.get("SKIP_INIT_DB"):
    try:
        init_db()
    except Exception as e:
        logger.warning("Database init warning (may be expected on Vercel): %s", e)

# Per-request DB connection reuse — see database.get_connection() docstring.
from database import close_request_connection  # noqa: E402
app.teardown_appcontext(close_request_connection)

# ---------------------------------------------------------------------------
# Background scraper state
# ---------------------------------------------------------------------------

scraper_status = {
    "running": False,
    "phase": "idle",
    "portal_progress": {},
    "done_portals": 0,
    "total_portals": 0,
    "total_jobs": 0,
    "qualified_jobs": 0,
    "inserted": 0,
    "skipped": 0,
    "digest_path": None,
    "error": None,
    "started_at": None,
    "finished_at": None,
}
scraper_lock = threading.Lock()
_scraper_stop_event = threading.Event()        # stop requested from UI "Stop" button
_scheduled_stop_event = threading.Event()      # stop requested for scheduled runs only
_is_scheduled_run = False                      # True when current run was triggered by scheduler

# ---------------------------------------------------------------------------
# AI agent run state
# ---------------------------------------------------------------------------
agent_status = {"running": False, "queued": 0, "error": None, "finished_at": None}
agent_lock = threading.Lock()

# ---------------------------------------------------------------------------
# Live search state
# ---------------------------------------------------------------------------

live_search_status = {
    "running": False,
    "phase": "idle",
    "portal_progress": {},
    "done_portals": 0,
    "total_portals": 0,
    "total_jobs": 0,
    "qualified_jobs": 0,
    "inserted": 0,
    "skipped": 0,
    "error": None,
    "started_at": None,
    "finished_at": None,
    "result_job_ids": [],
}
live_search_lock = threading.Lock()

# ---------------------------------------------------------------------------
# Daily scheduler (11:00 AM)
# ---------------------------------------------------------------------------

_scheduler = None


def _scheduled_pipeline_run():
    """Callback for the daily scheduled scraper run."""
    global scraper_status, _is_scheduled_run
    with scraper_lock:
        if scraper_status["running"]:
            logger.info("Scheduled run skipped - scraper is already running")
            return
        scraper_status = {
            "running": True,
            "phase": "starting",
            "portal_progress": {},
            "done_portals": 0,
            "total_portals": 0,
            "total_jobs": 0,
            "qualified_jobs": 0,
            "inserted": 0,
            "skipped": 0,
            "digest_path": None,
            "error": None,
            "started_at": datetime.now().isoformat(),
            "finished_at": None,
        }
    _scraper_stop_event.clear()
    _scheduled_stop_event.clear()
    _is_scheduled_run = True
    logger.info("Scheduled daily pipeline run starting")
    _run_scraper_pipeline()

    # Run AI agent pipeline after scraping
    try:
        from agent.graph import run_agent_pipeline
        import json as _json
        prefs = load_preferences() or DEFAULT_PREFS.copy()
        with open(os.path.join(BASE_DIR, "config.json")) as f:
            _config = _json.load(f)
        run_agent_pipeline(prefs, _config)
    except Exception as e:
        logger.error("AI agent pipeline error in scheduler: %s", e)


_HR_EMAIL_DIR = os.path.join(os.path.expanduser("~"), "Documents", "Claude")


def _load_hm():
    """Return the hiring_managers_search module (now bundled in the repo)."""
    import hiring_managers_search
    return hiring_managers_search


def _load_gmail():
    """Load send_gmail module from Documents/Claude via importlib."""
    import importlib.util as _ilu
    spec = _ilu.spec_from_file_location("send_gmail",
                                         os.path.join(_HR_EMAIL_DIR, "send_gmail.py"))
    mod = _ilu.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def setup_background_scheduler():
    """
    Start a simple time-checker scheduler. Replaces APScheduler to avoid
    executor-blocking issues caused by long-running scraper jobs.

    The scheduler loop wakes every 60 seconds, checks the current time, and
    fires jobs that haven't run yet today. Each job runs in its own daemon
    thread so the loop is never blocked.
    """
    global _scheduler

    # Keep APScheduler alive for the /api/scheduler/jobs endpoint compatibility
    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        from apscheduler.triggers.cron import CronTrigger
        _scheduler = BackgroundScheduler(daemon=True)
        # Display-only job stubs — real execution is in _start_simple_scheduler().
        # Callbacks are no-ops; these entries exist solely for the /api/scheduler/jobs UI.
        _scheduler.add_job(lambda: None, trigger=CronTrigger(hour=7, minute=0),
                           id="morning_pipeline",
                           name="Morning job scraper pipeline at 07:00",
                           replace_existing=True)
        _scheduler.add_job(lambda: None, trigger=CronTrigger(hour=19, minute=0),
                           id="evening_pipeline",
                           name="Evening job scraper pipeline at 19:00",
                           replace_existing=True)
        _scheduler.start()
    except Exception:
        pass  # API compatibility only — actual scheduling is done below

    _start_simple_scheduler()


def _start_simple_scheduler():
    """
    Reliable minute-tick scheduler that fires jobs in isolated daemon threads.
    Cannot be blocked by long-running jobs.
    """
    import threading
    import time as _time
    from datetime import datetime as _dt, date as _date

    MORNING_PIPELINE_HOUR = 7
    EVENING_PIPELINE_HOUR = 19

    _SCHED_STATE_FILE = os.path.join(BASE_DIR, "data", "scheduler_state.json")

    def _load_state() -> dict:
        import json as _j
        try:
            with open(_SCHED_STATE_FILE) as f:
                return _j.load(f)
        except Exception:
            return {}

    def _save_state(state: dict):
        import json as _j
        try:
            os.makedirs(os.path.dirname(_SCHED_STATE_FILE), exist_ok=True)
            with open(_SCHED_STATE_FILE, "w") as f:
                _j.dump(state, f)
        except Exception as e:
            logger.warning("Scheduler state save failed: %s", e)

    _ran_today: dict = _load_state()   # persisted: job_key -> date string of last run

    def _already_ran(key: str, today: str) -> bool:
        return _ran_today.get(key) == today

    def _mark_ran(key: str, today: str):
        _ran_today[key] = today
        _save_state(_ran_today)

    def _fire(name: str, fn, *args):
        def _run():
            try:
                fn(*args)
            except Exception as e:
                logger.error("Simple scheduler job '%s' failed: %s", name, e)
        t = threading.Thread(target=_run, daemon=True, name=f"job-{name}")
        t.start()

    def _loop():
        logger.info("Simple scheduler started — pipeline@07:00/19:00")

        while True:
            _time.sleep(60)
            try:
                now   = _dt.now()
                today = _date.today().isoformat()
                h     = now.hour

                # Morning pipeline at 07:00
                if h >= MORNING_PIPELINE_HOUR and h < EVENING_PIPELINE_HOUR and not _already_ran("morning_pipeline", today):
                    _mark_ran("morning_pipeline", today)
                    logger.info("Simple scheduler: firing morning pipeline")
                    _fire("morning_pipeline", _scheduled_pipeline_run)

                # Evening pipeline at 19:00
                if h >= EVENING_PIPELINE_HOUR and not _already_ran("evening_pipeline", today):
                    _mark_ran("evening_pipeline", today)
                    logger.info("Simple scheduler: firing evening pipeline")
                    _fire("evening_pipeline", _scheduled_pipeline_run)


            except Exception as e:
                logger.error("Simple scheduler loop error: %s", e)

    t = threading.Thread(target=_loop, daemon=True, name="simple-scheduler")
    t.start()


def _run_apollo_enrichment(job_ids):
    """Run contact enrichment for a list of job IDs using the contact scraper."""
    conn = get_connection()
    cursor = conn.cursor()
    placeholders = ",".join("?" for _ in job_ids)
    cursor.execute(
        f"SELECT job_id, company, job_description, apply_url FROM job_listings "
        f"WHERE job_id IN ({placeholders}) "
        f"AND (poster_email IS NULL OR poster_email = '')",
        job_ids,
    )
    rows = [dict(r) for r in cursor.fetchall()]
    conn.close()

    if not rows:
        return

    prefs = load_preferences() or {}
    linkedin_email = prefs.get("linkedin_email", "").strip()
    linkedin_password = prefs.get("linkedin_password", "").strip()

    contacts = enrich_jobs_with_contacts(
        rows,
        linkedin_email=linkedin_email or None,
        linkedin_password=linkedin_password or None,
    )
    for jid, info in contacts.items():
        update_job_contacts(
            jid,
            info.get("poster_name", ""),
            info.get("poster_email", ""),
            info.get("poster_phone", ""),
            info.get("poster_linkedin", ""),
        )
        # Also update the job description if LinkedIn JSON-LD returned one
        jd = info.get("jd_text", "")
        if jd:
            update_job_description(jid, jd)


def _run_scraper_pipeline():
    """Run the full pipeline in a background thread."""
    global scraper_status, _is_scheduled_run

    def _stopped():
        # Scheduled runs are only stopped by the dedicated scheduled-stop event.
        # Manual UI runs are stopped by the regular stop event.
        if _is_scheduled_run:
            return _scheduled_stop_event.is_set()
        return _scraper_stop_event.is_set()

    def _mark_stopped():
        global _is_scheduled_run
        with scraper_lock:
            scraper_status["phase"] = "stopped"
            scraper_status["finished_at"] = datetime.now().isoformat()
            scraper_status["running"] = False
            _is_scheduled_run = False
        logger.info("Scraper stopped")

    try:
        config = load_config()
        preferences = apply_env_overrides(load_preferences() or DEFAULT_PREFS.copy())
        job_titles = preferences.get("job_titles", DEFAULT_PREFS["job_titles"])
        locations = preferences.get("locations", DEFAULT_PREFS["locations"])
        top_n = preferences.get("top_jobs_per_digest", 10)

        # Phase 1: Scrape
        with scraper_lock:
            scraper_status["phase"] = "scraping"
            scraper_status["portal_progress"] = {}

        def scrape_cb(portal, status, count, done, total):
            with scraper_lock:
                scraper_status["portal_progress"][portal] = {
                    "status": status, "count": count,
                }
                scraper_status["done_portals"] = done
                scraper_status["total_portals"] = total

        all_jobs, portal_results = scrape_all_portals(
            job_titles, locations, config, progress_callback=scrape_cb,
            stop_event=_scraper_stop_event,
        )

        if _stopped():
            _mark_stopped()
            return

        with scraper_lock:
            scraper_status["total_jobs"] = len(all_jobs)

        if not all_jobs:
            with scraper_lock:
                scraper_status["phase"] = "done"
                scraper_status["finished_at"] = datetime.now().isoformat()
                scraper_status["running"] = False
            return

        # Phase 2: Analyze
        if _stopped():
            _mark_stopped()
            return
        with scraper_lock:
            scraper_status["phase"] = "analyzing"

        qualified_jobs, all_analyzed = analyze_jobs(all_jobs, preferences, config)

        with scraper_lock:
            scraper_status["qualified_jobs"] = len(qualified_jobs)

        # Phase 3: Store
        with scraper_lock:
            scraper_status["phase"] = "storing"

        # Per-user CV scoring happens lazily on the read side via
        # _score_unscored_for_user; no pre-scoring needed here.
        for job in all_analyzed:
            job["job_id"] = generate_job_id(
                job["portal"], job["company"], job["role"], job.get("location", ""),
            )
        inserted, skipped = insert_jobs_bulk(all_analyzed)

        with scraper_lock:
            scraper_status["inserted"] = inserted
            scraper_status["skipped"] = skipped

        # Phase 3.5: Telegram alerts
        tg_token = preferences.get("telegram_bot_token", "").strip()
        tg_chat = preferences.get("telegram_chat_id", "").strip()
        tg_min = int(preferences.get("telegram_min_score", 65))
        if tg_token and tg_chat:
            with scraper_lock:
                scraper_status["phase"] = "telegram_alerts"
            alert_count = 0
            for job in qualified_jobs:
                if job.get("relevance_score", 0) >= tg_min:
                    send_telegram_alert(job, tg_token, tg_chat)
                    alert_count += 1
            if alert_count > 0 or inserted > 0:
                send_telegram_batch_summary(len(all_jobs), len(qualified_jobs), inserted, tg_token, tg_chat)
            logger.info("Sent %d Telegram alerts", alert_count)

        # Phase 3.6: Contact enrichment (via scraper, no API key needed)
        with scraper_lock:
            scraper_status["phase"] = "enriching_contacts"
        all_job_ids = [j["job_id"] for j in all_analyzed if j.get("job_id")]
        if all_job_ids:
            _run_apollo_enrichment(all_job_ids)

        # Phase 4: Digest
        with scraper_lock:
            scraper_status["phase"] = "generating_digest"

        # Prefer jobs not emailed in the last 7 days, but backfill so the digest
        # is never empty when few jobs qualify (e.g. LLM down -> keyword-only).
        digest_jobs = select_digest_jobs(qualified_jobs, top_n, days=7)
        stats = get_comprehensive_stats()
        html_path, _ = generate_digest(
            digest_jobs, portal_results, preferences, stats, open_browser=False,
        )

        sent_ids = [j.get("job_id") for j in digest_jobs if j.get("job_id")]
        if sent_ids:
            mark_sent_in_digest(sent_ids)

        with scraper_lock:
            scraper_status["digest_path"] = os.path.basename(html_path)

        # Phase 5: Email notification
        recipient = preferences.get("email", "").strip()
        gmail_addr = preferences.get("gmail_address", "").strip()
        gmail_pass = preferences.get("gmail_app_password", "").strip()
        if recipient and gmail_addr and gmail_pass:
            with scraper_lock:
                scraper_status["phase"] = "sending_email"
            try:
                email_jobs = digest_jobs if digest_jobs else []
                ok, _err = send_job_email(recipient, email_jobs, preferences)
                if ok:
                    logger.info("Email digest sent to %s", recipient)
                else:
                    logger.error("Email digest failed: %s", _err)
            except Exception as e:
                logger.error("Failed to send email: %s", e)

        # Phase 5.5: Cleanup old jobs (>30 days)
        with scraper_lock:
            scraper_status["phase"] = "cleanup"
        try:
            removed = delete_old_jobs(days=30)
            logger.info("Cleanup: removed %d jobs older than 30 days", removed)
        except Exception as e:
            logger.error("Cleanup error: %s", e)

        with scraper_lock:
            scraper_status["phase"] = "done"
            scraper_status["finished_at"] = datetime.now().isoformat()
            scraper_status["running"] = False

    except Exception as e:
        logger.exception("Scraper pipeline error")
        with scraper_lock:
            scraper_status["error"] = str(e)
            scraper_status["phase"] = "error"
            scraper_status["running"] = False


def _run_live_search(query, location):
    """Run a slim scrape+analyze+store pipeline for live search from the jobs page."""
    global live_search_status
    try:
        config = load_config()
        preferences = apply_env_overrides(load_preferences() or DEFAULT_PREFS.copy())

        job_titles = [query] if query else preferences.get("job_titles", DEFAULT_PREFS["job_titles"])
        locations_list = [location] if location else preferences.get("locations", DEFAULT_PREFS["locations"])

        # Phase 1: Scrape
        with live_search_lock:
            live_search_status["phase"] = "scraping"
            live_search_status["portal_progress"] = {}

        def scrape_cb(portal, status, count, done, total):
            with live_search_lock:
                live_search_status["portal_progress"][portal] = {
                    "status": status, "count": count,
                }
                live_search_status["done_portals"] = done
                live_search_status["total_portals"] = total

        all_jobs, portal_results = scrape_all_portals(
            job_titles, locations_list, config, progress_callback=scrape_cb,
        )

        with live_search_lock:
            live_search_status["total_jobs"] = len(all_jobs)

        if not all_jobs:
            with live_search_lock:
                live_search_status["phase"] = "done"
                live_search_status["finished_at"] = datetime.now().isoformat()
                live_search_status["running"] = False
            return

        # Phase 2: Analyze
        with live_search_lock:
            live_search_status["phase"] = "analyzing"

        qualified_jobs, all_analyzed = analyze_jobs(all_jobs, preferences, config)

        with live_search_lock:
            live_search_status["qualified_jobs"] = len(qualified_jobs)

        # Phase 3: Store
        with live_search_lock:
            live_search_status["phase"] = "storing"

        # Per-user CV scoring happens lazily on the read side; no
        # pre-scoring against any single user's CV here.
        for job in all_analyzed:
            job["job_id"] = generate_job_id(
                job["portal"], job["company"], job["role"], job.get("location", ""),
            )
        inserted, skipped = insert_jobs_bulk(all_analyzed)
        result_ids = [j["job_id"] for j in all_analyzed if j.get("job_id")]

        with live_search_lock:
            live_search_status["inserted"] = inserted
            live_search_status["skipped"] = skipped
            live_search_status["result_job_ids"] = result_ids

        # Phase 3.5: Telegram alerts
        tg_token = preferences.get("telegram_bot_token", "").strip()
        tg_chat = preferences.get("telegram_chat_id", "").strip()
        tg_min = int(preferences.get("telegram_min_score", 65))
        if tg_token and tg_chat:
            with live_search_lock:
                live_search_status["phase"] = "telegram_alerts"
            for job in qualified_jobs:
                if job.get("relevance_score", 0) >= tg_min:
                    send_telegram_alert(job, tg_token, tg_chat)

        # Phase 4: Contact enrichment (via scraper, no API key needed)
        if result_ids:
            with live_search_lock:
                live_search_status["phase"] = "enriching_contacts"
            _run_apollo_enrichment(result_ids)

        with live_search_lock:
            live_search_status["phase"] = "done"
            live_search_status["finished_at"] = datetime.now().isoformat()
            live_search_status["running"] = False

    except Exception as e:
        logger.exception("Live search pipeline error")
        with live_search_lock:
            live_search_status["error"] = str(e)
            live_search_status["phase"] = "error"
            live_search_status["running"] = False


# ---------------------------------------------------------------------------
# Scheduler & Telegram bot startup guards
# ---------------------------------------------------------------------------
# Skip on Vercel (serverless – no persistent processes).
# Flask dev mode: only start in the reloader child (WERKZEUG_RUN_MAIN=true).
# Gunicorn with --preload: code runs once in the arbiter, then workers fork.
#   The arbiter imports gunicorn.arbiter, workers do not – use that to detect
#   we are in a preloaded arbiter and start background tasks there (they
#   survive fork because they are daemon threads).
# Gunicorn without --preload: each worker imports app.py; with 1 worker this
#   is fine – start unconditionally when not in debug mode.


def _is_gunicorn_arbiter():
    """Return True if we are running inside the gunicorn arbiter (master)."""
    return "gunicorn" in os.environ.get("SERVER_SOFTWARE", "")


def _should_start_background_tasks():
    if _IS_VERCEL:
        return False
    # Flask dev server with reloader
    if app.debug:
        return os.environ.get("WERKZEUG_RUN_MAIN") == "true"
    # Gunicorn or any other production server
    return True


def _git_pull_and_sync():
    """
    Periodically git pull and import any new scrape committed by GitHub Actions.
    Runs every 30 minutes in a background thread.
    """
    import subprocess
    import time as _time

    while True:
        _time.sleep(1800)  # wait 30 minutes between checks
        try:
            result = subprocess.run(
                ["git", "pull", "--ff-only"],
                cwd=BASE_DIR,
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0 and "Already up to date." not in result.stdout:
                logger.info("git pull: %s", result.stdout.strip())
                sync_from_scrape(BASE_DIR, insert_jobs_bulk)
                # Auto-run agent pipeline so new jobs get scored and drafted immediately
                try:
                    import json as _json
                    from agent.graph import run_agent_pipeline
                    _prefs = load_preferences() or DEFAULT_PREFS.copy()
                    with open(os.path.join(BASE_DIR, "config.json")) as _f:
                        _cfg = _json.load(_f)
                    run_agent_pipeline(_prefs, _cfg)
                    logger.info("git sync: agent pipeline completed")
                except Exception as _ae:
                    logger.warning("git sync: agent pipeline error: %s", _ae)
            else:
                logger.debug("git pull: no new commits")
        except Exception as e:
            logger.warning("git pull sync failed: %s", e)


if _should_start_background_tasks():
    setup_background_scheduler()

    # Auto-import any scraped jobs committed by GitHub Actions (on startup)
    import threading as _threading
    _threading.Thread(
        target=sync_from_scrape,
        args=(BASE_DIR, insert_jobs_bulk),
        daemon=True,
    ).start()

    # Periodically pull latest scrape from GitHub Actions and import
    _threading.Thread(target=_git_pull_and_sync, daemon=True).start()

    # Start Telegram bot if token is configured
    _bot_prefs = apply_env_overrides(load_preferences() or DEFAULT_PREFS.copy())
    _bot_token = _bot_prefs.get("telegram_bot_token", "").strip()
    if _bot_token:
        start_telegram_bot(_bot_token)

# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Auth Routes
# ---------------------------------------------------------------------------

@app.route('/login')
def login():
    if session.get('user'):
        return redirect(url_for('dashboard'))
    return render_template('login.html', is_vercel=_IS_VERCEL,
                           enable_dev_login=_ENABLE_DEV_LOGIN,
                           enable_demo_login=_ENABLE_DEMO_LOGIN)

@app.route('/auth/google')
def auth_google():
    if not os.environ.get("GOOGLE_CLIENT_ID"):
        flash("Google OAuth is not configured.", "error")
        return redirect(url_for('login'))
    redirect_uri = url_for('auth_callback', _external=True)
    return google.authorize_redirect(redirect_uri)

def _signup_allowed(email: str) -> bool:
    """Open by default. Set ALLOWED_EMAILS (comma-separated) to restrict
    sign-ups to a specific list. Existing users are always grandfathered.
    """
    if not email:
        return False
    raw = os.environ.get("ALLOWED_EMAILS", "").strip()
    if not raw:
        return True  # open registration
    allowed = {e.strip().lower() for e in raw.split(",") if e.strip()}
    if email.lower() in allowed:
        return True
    # Grandfather any user who already has a row
    try:
        existing = get_user_by_email(email)
        return existing is not None
    except Exception:
        return False


@app.route('/auth/callback')
@csrf.exempt
def auth_callback():
    try:
        token = google.authorize_access_token()
        userinfo = token.get('userinfo') or {}
        email = (userinfo.get('email') or '').strip().lower()
        if not email:
            flash("Authentication failed: no email returned.", "error")
            return redirect(url_for('login'))
        if not _signup_allowed(email):
            logger.warning("Sign-in blocked by allowlist: %s", email)
            flash("This deployment is invite-only. Contact the owner to be added.", "error")
            return redirect(url_for('login'))
        uid = get_or_create_user(email, userinfo.get('name', ''), userinfo.get('picture', ''))
        # First user signing into a fresh deployment becomes the admin.
        try:
            promote_first_user_to_admin()
        except Exception as e:
            logger.warning("First-user admin bootstrap failed: %s", e)
        session.permanent = True
        session['user'] = {
            'email': email,
            'name':  userinfo.get('name', ''),
            'picture': userinfo.get('picture', ''),
            'id':    uid,
        }
    except Exception as e:
        logger.error("OAuth callback error: %s", e)
        flash("Authentication failed.", "error")
        return redirect(url_for('login'))
    next_url = request.args.get('next') or ''
    # Reject external redirects (open-redirect fix)
    from urllib.parse import urlparse
    if urlparse(next_url).netloc or next_url.startswith('//'):
        next_url = ''
    # New users (no saved preferences) go to preferences page with welcome flag
    if not next_url:
        try:
            existing_prefs = load_preferences(uid) or {}
            has_setup = bool(existing_prefs.get("job_titles") and existing_prefs.get("locations"))
        except Exception:
            has_setup = True
        if not has_setup:
            next_url = url_for('preferences') + '?welcome=1'
        else:
            next_url = url_for('dashboard')
    return redirect(next_url)

@app.route('/auth/dev-login', methods=['POST'])
@csrf.exempt
def auth_dev_login():
    """Local dev only — gated by ENABLE_DEV_LOGIN=1 (and never on Vercel)."""
    if not _ENABLE_DEV_LOGIN:
        from flask import abort
        abort(403)
    uid = get_or_create_user('dev@localhost', 'Dev User')
    # First-time bootstrap: dev user becomes admin so local dev "just works"
    try:
        promote_first_user_to_admin()
    except Exception:
        pass
    session.permanent = True
    session['user'] = {'email': 'dev@localhost', 'name': 'Dev User', 'picture': '', 'id': uid}
    return redirect(url_for('dashboard'))


@app.route('/auth/demo-login', methods=['POST'])
@csrf.exempt
def auth_demo_login():
    """
    Public, no-Google demo entry for first-time-flow testing. Temporary:
    gated by ENABLE_DEMO_LOGIN (on by default; set to 0 to disable). Each call
    mints a FRESH throwaway user so the visitor always experiences the
    first-time flow (empty dashboard -> CV upload prompt). Demo users are never
    admins and are blocked from costly actions (see _is_demo_user).
    """
    if not _ENABLE_DEMO_LOGIN:
        from flask import abort
        abort(403)
    # Bound table growth: clear out throwaway demo users older than 24h before
    # minting a new one, so the public endpoint can't grow the users table
    # without limit.
    try:
        purge_stale_demo_users(_DEMO_EMAIL_DOMAIN, hours=24)
    except Exception as e:
        logger.warning("purge_stale_demo_users failed: %s", e)
    import secrets as _secrets
    token = _secrets.token_hex(4)
    email = f"demo-{token}@{_DEMO_EMAIL_DOMAIN}"
    uid = get_or_create_user(email, 'Demo User')
    # Deliberately NOT promoted to admin — demo users stay least-privileged.
    session.permanent = True
    session['user'] = {'email': email, 'name': 'Demo User', 'picture': '', 'id': uid, 'demo': True}
    logger.info("Demo session started: %s", email)
    return redirect(url_for('dashboard'))

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.context_processor
def inject_globals():
    """Make a few values available to every template (footer year, etc.)."""
    return {
        'current_year': datetime.now().year,
        'enable_demo_login': _ENABLE_DEMO_LOGIN,
        'is_demo_user': _is_demo_user(),
    }


@app.route("/")
def index():
    # Authenticated users go straight to their workspace; everyone else sees
    # the public landing page that explains what the product does.
    if session.get('user'):
        return redirect(url_for('dashboard'))
    return render_template('landing.html')


@app.route("/privacy")
def privacy():
    return render_template('privacy.html')


@app.route("/terms")
def terms():
    return render_template('terms.html')


@app.route("/favicon.ico")
def favicon():
    # Return 204 No Content to avoid 500 errors when favicon is missing
    return "", 204


@app.route("/dashboard")
def dashboard():
    """
    Trimmed dashboard: three stats (scraped today, applied today,
    high-match ≥80% today) and the top 10 best-matched jobs from the
    last 24 hours with apply links. Everything else removed.
    """
    uid = current_user_id()
    cv_data = load_cv_data(uid) if uid else None
    cv_uploaded = cv_data is not None
    _sess_user = dict(session).get('user', {})
    user_email = _sess_user.get('email', '')
    # First name for the greeting — fall back to the email local-part.
    _name = (_sess_user.get('name') or '').strip()
    user_name = (_name.split()[0] if _name else (user_email.split('@')[0] if user_email else ''))

    # Score any new jobs before querying — raise limit so a fresh scrape batch
    # (typically 500–1000 jobs) gets covered in one dashboard hit.
    if uid and cv_uploaded:
        try:
            _score_unscored_for_user(uid, limit=1500)
        except Exception as e:
            logger.warning("Dashboard lazy scoring failed for user %s: %s", uid, e)

    # Always use UTC so the cutoff matches the UTC-stored date_found values.
    from datetime import timezone as _tz
    _now_utc = datetime.now(_tz.utc).replace(tzinfo=None)
    today_str = _now_utc.strftime("%Y-%m-%d")

    conn = get_connection()
    cur = conn.cursor()

    # 1) Jobs scraped today (shared)
    cur.execute(
        "SELECT COUNT(*) AS n FROM job_listings WHERE date_found LIKE ?",
        (f"{today_str}%",),
    )
    jobs_today = cur.fetchone()["n"]

    # 2) Jobs the user viewed today
    viewed_today = 0
    if uid:
        cur.execute(
            """
            SELECT COUNT(*) AS n FROM user_job_state
            WHERE user_id = ? AND viewed_at LIKE ?
            """,
            (uid, f"{today_str}%"),
        )
        viewed_today = cur.fetchone()["n"]

    # 3) High-match (≥80) jobs found today, scored for this user
    high_match_today = 0
    if uid:
        cur.execute(
            """
            SELECT COUNT(*) AS n
            FROM user_job_state s
            JOIN job_listings j ON j.job_id = s.job_id
            WHERE s.user_id = ?
              AND j.date_found LIKE ?
              AND s.cv_score >= 80
            """,
            (uid, f"{today_str}%"),
        )
        high_match_today = cur.fetchone()["n"]

    # 4) Top 10 best-matched jobs from the last 24 hours strictly.
    #    24-hour window is the filter; cv_score is the sort. No fallback.
    top_jobs = []
    if uid and cv_uploaded:
        _batch_cutoff = (datetime.now(_tz.utc).replace(tzinfo=None) - timedelta(hours=24)).isoformat()
        cur.execute(
            """
            SELECT j.job_id, j.role, j.company, j.location,
                   j.remote_status, j.salary, j.portal, j.apply_url,
                   j.date_found, j.date_first_seen,
                   COALESCE(s.cv_score, 0) AS cv_score,
                   s.viewed_at
            FROM job_listings j
            LEFT JOIN user_job_state s
              ON s.job_id = j.job_id AND s.user_id = ?
            WHERE COALESCE(j.date_first_seen, j.date_found) >= ?
              AND COALESCE(s.hidden, 0) = 0
              AND COALESCE(s.applied_status, 0) = 0
            ORDER BY COALESCE(s.cv_score, 0) DESC, j.date_found DESC
            LIMIT 10
            """,
            (uid, _batch_cutoff),
        )
        top_jobs = [dict(r) for r in cur.fetchall()]
    conn.close()

    return render_template(
        "dashboard.html",
        jobs_today=jobs_today,
        viewed_today=viewed_today,
        high_match_today=high_match_today,
        top_jobs=top_jobs,
        cv_uploaded=cv_uploaded,
        user_email=user_email,
        user_name=user_name,
    )


@app.route("/api/dashboard/top-jobs")
def dashboard_top_jobs():
    """Return top 10 jobs by CV score (or relevance score if no CV)."""
    uid = current_user_id()
    cv_data = load_cv_data(uid) if uid else None
    conn = get_connection()
    c = conn.cursor()
    if uid and cv_data:
        c.execute("""
            SELECT j.job_id, j.role, j.company, j.location, j.relevance_score,
                   COALESCE(s.cv_score, 0) AS cv_score,
                   j.apply_url, j.date_found, j.portal, j.remote_status
            FROM job_listings j
            LEFT JOIN user_job_state s ON s.job_id = j.job_id AND s.user_id = ?
            WHERE COALESCE(s.cv_score, 0) > 0
              AND COALESCE(s.hidden, 0) = 0
            ORDER BY COALESCE(s.cv_score, 0) DESC, j.relevance_score DESC
            LIMIT 10
        """, (uid,))
    elif uid:
        c.execute("""
            SELECT j.job_id, j.role, j.company, j.location, j.relevance_score,
                   COALESCE(s.cv_score, 0) AS cv_score,
                   j.apply_url, j.date_found, j.portal, j.remote_status
            FROM job_listings j
            LEFT JOIN user_job_state s ON s.job_id = j.job_id AND s.user_id = ?
            WHERE COALESCE(s.hidden, 0) = 0
            ORDER BY j.relevance_score DESC, j.date_found DESC
            LIMIT 10
        """, (uid,))
    elif cv_data:
        c.execute("""
            SELECT job_id, role, company, location, relevance_score, cv_score,
                   apply_url, date_found, portal, remote_status
            FROM job_listings
            WHERE cv_score > 0 AND (hidden=0 OR hidden IS NULL)
            ORDER BY cv_score DESC, relevance_score DESC
            LIMIT 10
        """)
    else:
        c.execute("""
            SELECT job_id, role, company, location, relevance_score, cv_score,
                   apply_url, date_found, portal, remote_status
            FROM job_listings
            WHERE (hidden=0 OR hidden IS NULL)
            ORDER BY relevance_score DESC, date_found DESC
            LIMIT 10
        """)
    jobs = [dict(r) for r in c.fetchall()]
    conn.close()
    return jsonify({"jobs": jobs, "cv_uploaded": cv_data is not None})


def _build_jobs_query(filters, user_id: int = None):
    """
    Build SQL WHERE clause and params from a filters dict.
    Returns (conditions, params, order, join_sql, join_params, select_extra).
    Callers compose:  SELECT <cols>, <select_extra> FROM job_listings <join_sql>
                      WHERE <AND of conditions> ORDER BY <order>
    `user_id` scopes per-user columns (applied_status / hidden / cv_score /
    notes) via a LEFT JOIN on user_job_state.
    """
    conditions = []
    params = []
    select_extra = ""

    search = filters.get("search", "")
    portal = filters.get("portal", "")
    remote = filters.get("remote", "")
    company_type = filters.get("company_type", "")
    sort = filters.get("sort", "date_desc")
    applied = filters.get("applied", "")
    location = filters.get("location", "")
    recency = filters.get("recency", "")
    min_score = filters.get("min_score", "0")
    experience = filters.get("experience", "")
    salary_min = filters.get("salary_min", "")
    salary_max = filters.get("salary_max", "")
    company_stage = filters.get("company_stage", "")

    join_sql = ""
    join_params: list = []
    if user_id:
        join_sql = (
            " LEFT JOIN user_job_state ujs "
            "ON ujs.job_id = job_listings.job_id AND ujs.user_id = ?"
        )
        join_params = [int(user_id)]
        select_extra = (
            ", COALESCE(ujs.applied_status, 0) AS user_applied_status, "
            "COALESCE(ujs.cv_score, 0) AS user_cv_score, "
            "ujs.user_notes AS user_notes, "
            "ujs.applied_date AS user_applied_date, "
            "ujs.follow_up_date AS user_follow_up_date, "
            "ujs.rejection_reason AS user_rejection_reason, "
            "COALESCE(ujs.hidden, 0) AS user_hidden"
        )

    # Always exclude hidden jobs (hidden = 1) unless explicitly requested.
    # When user_id is in scope we honour the per-user hidden flag; otherwise
    # fall back to the legacy column on job_listings.
    if not filters.get("show_hidden"):
        if user_id:
            conditions.append("(COALESCE(ujs.hidden, 0) = 0)")
        else:
            conditions.append("(hidden = 0 OR hidden IS NULL)")

    # Exclude international locations by default unless a specific location is
    # chosen (in which case the user knows what they're filtering to) or
    # show_international is explicitly set.
    if not filters.get("location") and not filters.get("show_international"):
        intl = list(_INTERNATIONAL_CANONICALS)
        intl_ph = ",".join("?" for _ in intl)
        kw_conds = " OR ".join("LOWER(location) LIKE ?" for _ in _INTERNATIONAL_KEYWORDS)
        conditions.append(
            f"(location IS NULL OR location = '' OR "
            f"(location NOT IN ({intl_ph}) AND NOT ({kw_conds})))"
        )
        params.extend(intl)
        params.extend(f"%{kw}%" for kw in _INTERNATIONAL_KEYWORDS)

    # Default minimum score filter (0 = show all)
    try:
        min_score_val = int(min_score)
    except (ValueError, TypeError):
        min_score_val = 0
    if min_score_val > 0:
        conditions.append("relevance_score >= ?")
        params.append(min_score_val)

    if search:
        conditions.append("(role LIKE ? OR company LIKE ? OR job_description LIKE ?)")
        like = f"%{search}%"
        params.extend([like, like, like])
    if portal:
        conditions.append("portal = ?")
        params.append(portal)
    if remote:
        conditions.append("remote_status = ?")
        params.append(remote)
    if company_type:
        conditions.append("company_type = ?")
        params.append(company_type)
    if location:
        if location == "Others":
            # Exclude the 4 pinned cities; show remaining non-null locations
            _pinned = ["Pune", "Mumbai", "Bangalore", "Hyderabad"]
            excl_patterns = []
            for city in _pinned:
                excl_patterns.extend(_CITY_PATTERNS.get(city, []))
            not_clauses = " AND ".join("location NOT LIKE ?" for _ in excl_patterns)
            conditions.append(
                f"(location IS NOT NULL AND location != '' AND ({not_clauses}))"
            )
            params.extend([f"%{p}%" for p in excl_patterns])
        else:
            city_patterns = _CITY_PATTERNS.get(location)
            if city_patterns:
                like_clauses = ["location LIKE ?" for _ in city_patterns]
                conditions.append("(" + " OR ".join(like_clauses) + ")")
                params.extend([f"%{p}%" for p in city_patterns])
            else:
                conditions.append("location LIKE ?")
                params.append(f"%{location}%")
    if recency:
        recency_map = {
            "24h": timedelta(hours=24),
            "3d": timedelta(days=3),
            "1w": timedelta(weeks=1),
            "1m": timedelta(days=30),
        }
        td = recency_map.get(recency)
        if td:
            cutoff_date = (datetime.now() - td).strftime("%Y-%m-%d")
            conditions.append(
                "(date_posted IS NOT NULL AND date_posted != '' AND date_posted >= ?)"
            )
            params.append(cutoff_date)

    if user_id:
        applied_map = {
            "none": "COALESCE(ujs.applied_status, 0) = 0",
            "applied": "ujs.applied_status = 1",
            "saved": "ujs.applied_status = 2",
            "phone_screen": "ujs.applied_status = 3",
            "interview": "ujs.applied_status = 4",
            "offer": "ujs.applied_status = 5",
            "rejected": "ujs.applied_status = 6",
        }
    else:
        applied_map = {
            "none": "applied_status = 0",
            "applied": "applied_status = 1",
            "saved": "applied_status = 2",
            "phone_screen": "applied_status = 3",
            "interview": "applied_status = 4",
            "offer": "applied_status = 5",
            "rejected": "applied_status = 6",
        }
    if applied in applied_map:
        conditions.append(applied_map[applied])

    if experience:
        exp_ranges = {
            "0-3": (0, 3),
            "3-7": (3, 7),
            "7-12": (7, 12),
            "12+": (12, 99),
        }
        exp_range = exp_ranges.get(experience)
        if exp_range:
            lo, hi = exp_range
            conditions.append(
                "(experience_min IS NOT NULL AND experience_min <= ? AND experience_max >= ?)"
            )
            params.extend([hi, lo])

    if salary_min:
        try:
            sal_min_inr = int(salary_min) * 100_000
            conditions.append("(salary_min IS NOT NULL AND salary_max >= ?)")
            params.append(sal_min_inr)
        except (ValueError, TypeError):
            pass
    if salary_max:
        try:
            sal_max_inr = int(salary_max) * 100_000
            conditions.append("(salary_min IS NOT NULL AND salary_min <= ?)")
            params.append(sal_max_inr)
        except (ValueError, TypeError):
            pass

    if company_stage:
        conditions.append("company_funding_stage = ?")
        params.append(company_stage)

    if user_id:
        sort_map = {
            "score_desc": "relevance_score DESC, date_found DESC",
            "score_asc": "relevance_score ASC, date_found DESC",
            "date_desc": "date_found DESC",
            "date_asc": "date_found ASC",
            "company_asc": "company ASC",
            "cv_score_desc": "COALESCE(ujs.cv_score, 0) DESC, relevance_score DESC, date_found DESC",
            "cv_score_asc": "COALESCE(ujs.cv_score, 0) ASC, relevance_score ASC, date_found DESC",
        }
    else:
        sort_map = {
            "score_desc": "relevance_score DESC, date_found DESC",
            "score_asc": "relevance_score ASC, date_found DESC",
            "date_desc": "date_found DESC",
            "date_asc": "date_found ASC",
            "company_asc": "company ASC",
            "cv_score_desc": "cv_score DESC, relevance_score DESC, date_found DESC",
            "cv_score_asc": "cv_score ASC, relevance_score ASC, date_found DESC",
        }
    # When CV is uploaded, "score" sorts should use cv_score (the displayed score)
    if filters.get("cv_uploaded") and sort in ("score_desc", "score_asc"):
        sort = "cv_score_desc" if sort == "score_desc" else "cv_score_asc"
    order = sort_map.get(sort, "date_found DESC")

    return conditions, params, order, join_sql, join_params, select_extra


@app.route("/jobs")
def jobs():
    uid = current_user_id()
    cv_data = load_cv_data(uid) if uid else None
    cv_uploaded = cv_data is not None

    # Read filter params
    filters = {
        "search": request.args.get("search", "").strip(),
        "portal": request.args.get("portal", ""),
        "remote": request.args.get("remote", ""),
        "company_type": request.args.get("company_type", ""),
        "sort": request.args.get("sort", "date_desc"),
        "applied": request.args.get("applied", ""),
        "location": request.args.get("location", ""),
        "recency": request.args.get("recency", ""),
        "min_score": request.args.get("min_score", "0"),
        "experience": request.args.get("experience", ""),
        "salary_min": request.args.get("salary_min", ""),
        "salary_max": request.args.get("salary_max", ""),
        "company_stage": request.args.get("company_stage", ""),
        "cv_uploaded": cv_uploaded,
    }
    conditions, params, order, join_sql, join_params, select_extra = _build_jobs_query(filters, user_id=uid)
    where = " WHERE " + " AND ".join(conditions) if conditions else ""

    conn = get_connection()
    cursor = conn.cursor()

    # Pagination params
    try:
        page = int(request.args.get("page", 1))
    except ValueError:
        page = 1
    per_page = 25
    offset = (page - 1) * per_page

    # Fetch total count
    cursor.execute(
        f"SELECT COUNT(*) as count FROM job_listings{join_sql}{where}",
        join_params + params,
    )
    total = cursor.fetchone()["count"]

    # Fetch paginated matching jobs
    cursor.execute(
        f"SELECT job_listings.*{select_extra} FROM job_listings{join_sql}{where} "
        f"ORDER BY {order} LIMIT ? OFFSET ?",
        join_params + params + [per_page, offset],
    )
    rows = [dict(r) for r in cursor.fetchall()]
    conn.close()
    # Overlay per-user state onto the legacy columns so templates / clients
    # that still read job.applied_status / cv_score / user_notes keep working.
    if uid:
        for r in rows:
            if "user_applied_status" in r:
                r["applied_status"] = r.pop("user_applied_status")
            if "user_cv_score" in r:
                r["cv_score"] = r.pop("user_cv_score")
            if "user_notes" in r and r["user_notes"] is not None:
                pass  # column name already matches
            if "user_applied_date" in r:
                r["applied_date"] = r.pop("user_applied_date")
            if "user_follow_up_date" in r:
                r["follow_up_date"] = r.pop("user_follow_up_date")
            if "user_rejection_reason" in r:
                r["rejection_reason"] = r.pop("user_rejection_reason")
            if "user_hidden" in r:
                r["hidden"] = r.pop("user_hidden")

    # Attach inline gap data if CV is uploaded
    if cv_uploaded:
        # Dynamic missing skills analysis ONLY (score already pre-computed or we just rely on cv_score)
        for job in rows:
            gap = compute_gap_analysis(job, cv_data)
            job["_missing_top3"] = gap.get("missing_skills", [])[:3]
            # Since we're paginated, we rely on the db column `cv_score` for sorting
            # But let's also pass the computed score incase it changed
            job["_cv_score"] = gap.get("cv_score", job.get("cv_score", 0))
    else:
        for job in rows:
            job["_missing_top3"] = []
            job["_cv_score"] = 0

    if request.headers.get("X-Requested-With") == "XMLHttpRequest" or request.args.get("ajax") == "1":
        return render_template("_job_card_list.html", jobs=rows, cv_uploaded=cv_uploaded, offset=offset)

    # Get distinct portals for filter dropdown
    conn2 = get_connection()
    cur2 = conn2.cursor()
    cur2.execute("SELECT DISTINCT portal FROM job_listings ORDER BY portal")
    portals = [r["portal"] for r in cur2.fetchall()]
    conn2.close()

    _internal_keys = {"cv_uploaded", "show_hidden", "show_international"}
    clean_filters = {k: v for k, v in filters.items() if v and v != "0" and k not in _internal_keys}

    prefs = (load_preferences(uid) if uid else None) or DEFAULT_PREFS.copy()

    return render_template(
        "jobs.html",
        jobs=rows, total=total,
        portals=portals,
        filters=filters, clean_filters=clean_filters,
        cv_uploaded=cv_uploaded,
        prefs=prefs,
    )


@app.route("/api/nlp-search", methods=["POST"])
def nlp_search():
    data = request.get_json(silent=True) or {}
    query = (data.get("query") or "").strip()
    if not query:
        return jsonify({"ok": False, "error": "Empty query"}), 400

    config = load_config()
    filters = parse_nlp_query(query, config)

    # Default min_score to 0 for NLP search (show all matching jobs)
    if "min_score" not in filters:
        filters["min_score"] = "0"

    uid = current_user_id()
    conditions, params, order, join_sql, join_params, select_extra = _build_jobs_query(filters, user_id=uid)
    where = " WHERE " + " AND ".join(conditions) if conditions else ""

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute(
        f"SELECT COUNT(*) as cnt FROM job_listings{join_sql}{where}",
        join_params + params,
    )
    total = cursor.fetchone()["cnt"]

    cursor.execute(
        f"SELECT job_listings.*{select_extra} FROM job_listings{join_sql}{where} "
        f"ORDER BY {order} LIMIT 25",
        join_params + params,
    )
    rows = [dict(r) for r in cursor.fetchall()]
    conn.close()
    if uid:
        for r in rows:
            if "user_applied_status" in r:
                r["applied_status"] = r.pop("user_applied_status")
            if "user_cv_score" in r:
                r["cv_score"] = r.pop("user_cv_score")
            if "user_hidden" in r:
                r["hidden"] = r.pop("user_hidden")

    # Build human-readable filter descriptions
    filter_labels = []
    if filters.get("search"):
        filter_labels.append(f'Search: "{filters["search"]}"')
    if filters.get("location"):
        filter_labels.append(f'Location: {filters["location"]}')
    if filters.get("remote"):
        filter_labels.append(f'Remote: {filters["remote"].title()}')
    if filters.get("salary_min"):
        filter_labels.append(f'Salary > {filters["salary_min"]}L')
    if filters.get("salary_max"):
        filter_labels.append(f'Salary < {filters["salary_max"]}L')
    if filters.get("experience"):
        filter_labels.append(f'Experience: {filters["experience"]} yrs')
    if filters.get("company_type"):
        filter_labels.append(f'Company: {filters["company_type"].title()}')
    if filters.get("applied"):
        filter_labels.append(f'Status: {filters["applied"]}')

    return jsonify({
        "ok": True,
        "query": query,
        "filters": filters,
        "filter_labels": filter_labels,
        "jobs": rows,
        "total": total,
    })


@app.route("/api/jobs", methods=["GET"])
def api_jobs_list():
    """
    Lightweight JSON list of jobs used by the /jobs page client-side render.

    Returns the freshest visible jobs joined with the current user's per-user
    state (cv_score, applied_status, hidden) so the UI can show personalised
    badges without re-running the heavy SQL builder.

    Query params:
        limit         int  (default 1000, max 5000)
        min_score     int  (default 0)
        all_locations bool (default 0) — set to 1 to bypass preference filter
    """
    uid = current_user_id()
    try:
        limit = max(1, min(5000, int(request.args.get("limit", 1000))))
    except (ValueError, TypeError):
        limit = 1000
    try:
        min_score = max(0, min(100, int(request.args.get("min_score", 0))))
    except (ValueError, TypeError):
        min_score = 0

    # Lazy per-user scoring — any jobs the user hasn't seen scored
    # (e.g. fresh from this morning's scrape) get scored now against
    # THIS user's CV + preferences. No-op when user has no CV.
    if uid:
        try:
            _score_unscored_for_user(uid, limit=limit)
        except Exception as e:
            logger.warning("Lazy scoring failed for user %s: %s", uid, e)

    # Build location filter from user preferences (skip if all_locations=1)
    loc_where = ""
    loc_params: list = []
    all_locations = request.args.get("all_locations", "0") not in ("", "0", "false")
    user_prefs: dict = {}
    if uid and not all_locations:
        try:
            user_prefs = load_preferences(uid) or {}
        except Exception:
            user_prefs = {}
        pref_locs = [l.strip() for l in user_prefs.get("locations", []) if l.strip()]
        if pref_locs:
            loc_conds = []
            for loc in pref_locs:
                loc_lower = loc.lower()
                if loc_lower in ("remote", "wfh", "work from home", "anywhere"):
                    loc_conds.append("LOWER(COALESCE(j.remote_status, '')) LIKE '%remote%'")
                else:
                    patterns = _CITY_PATTERNS.get(loc, [loc_lower])
                    for p in patterns:
                        loc_conds.append("LOWER(j.location) LIKE ?")
                        loc_params.append(f"%{p}%")
            if loc_conds:
                loc_where = " AND (" + " OR ".join(loc_conds) + ")"

    # Work mode filter from user preferences
    mode_where = ""
    mode_params: list = []
    _all_modes = {"on-site", "remote", "hybrid"}
    if uid and not all_locations:
        try:
            work_modes = set(user_prefs.get("work_modes") or _all_modes)
        except Exception:
            work_modes = _all_modes
        if work_modes and work_modes != _all_modes:
            placeholders = ",".join("?" for _ in work_modes)
            mode_where = f" AND j.remote_status IN ({placeholders})"
            mode_params = list(work_modes)

    conn = get_connection()
    cursor = conn.cursor()
    if uid:
        cursor.execute(
            f"""
            SELECT j.job_id, j.role, j.company, j.location, j.salary,
                   j.salary_currency, j.remote_status, j.portal, j.apply_url,
                   j.job_description, j.relevance_score, j.date_found,
                   j.date_first_seen, j.date_posted,
                   j.experience_min, j.experience_max,
                   COALESCE(s.cv_score, 0)       AS cv_score,
                   COALESCE(s.applied_status, 0) AS applied_status,
                   s.applied_date                AS applied_date,
                   s.user_notes                  AS user_notes,
                   COALESCE(s.hidden, 0)         AS hidden
            FROM job_listings j
            LEFT JOIN user_job_state s
              ON s.job_id = j.job_id AND s.user_id = ?
            WHERE COALESCE(j.relevance_score, 0) >= ?
              AND COALESCE(s.hidden, 0) = 0
              {loc_where}{mode_where}
            ORDER BY COALESCE(j.date_first_seen, j.date_found) DESC
            LIMIT ?
            """,
            [uid, min_score] + loc_params + mode_params + [limit],
        )
    else:
        cursor.execute(
            """
            SELECT job_id, role, company, location, salary, salary_currency,
                   remote_status, portal, apply_url, job_description,
                   relevance_score, date_found, date_first_seen, date_posted,
                   experience_min, experience_max,
                   COALESCE(cv_score, 0) AS cv_score,
                   COALESCE(applied_status, 0) AS applied_status,
                   applied_date, user_notes,
                   COALESCE(hidden, 0) AS hidden
            FROM job_listings
            WHERE COALESCE(relevance_score, 0) >= ?
              AND (hidden = 0 OR hidden IS NULL)
            ORDER BY COALESCE(date_first_seen, date_found) DESC
            LIMIT ?
            """,
            (min_score, limit),
        )
    jobs = [dict(r) for r in cursor.fetchall()]
    conn.close()
    return jsonify({"jobs": jobs, "count": len(jobs), "location_filtered": bool(loc_where), "mode_filtered": bool(mode_where)})


@app.route("/api/jobs/<job_id>/status", methods=["POST"])
def update_job_status(job_id):
    uid = require_user_id()
    data = request.get_json(silent=True) or {}
    status = data.get("status", 0)
    notes = data.get("notes")
    follow_up_date = data.get("follow_up_date")
    rejection_reason = data.get("rejection_reason")
    try:
        status = int(status)
    except (ValueError, TypeError):
        status = 0
    update_applied_status_user(uid, job_id, status, notes, follow_up_date, rejection_reason)
    return jsonify({"ok": True, "job_id": job_id, "status": status})


@csrf.exempt
@app.route("/api/jobs/<job_id>/view", methods=["POST"])
def mark_job_viewed_route(job_id):
    uid = require_user_id()
    mark_job_viewed(uid, job_id)
    return jsonify({"ok": True})


@app.route("/api/jobs/<job_id>/hide", methods=["POST"])
def hide_job_route(job_id):
    """Hide or unhide a job. Body: {hidden: true/false}"""
    uid = require_user_id()
    data = request.get_json(silent=True) or {}
    is_hidden = bool(data.get("hidden", True))
    hide_job_user(uid, job_id, is_hidden)
    return jsonify({"ok": True, "job_id": job_id, "hidden": is_hidden})


@app.route("/api/jobs/<job_id>/notes", methods=["POST"])
def save_job_notes(job_id):
    """Save user notes for a job without changing other fields."""
    uid = require_user_id()
    data = request.get_json(silent=True) or {}
    notes = data.get("notes", "")
    update_job_notes_user(uid, job_id, notes)
    return jsonify({"ok": True, "job_id": job_id})


@app.route("/api/admin/dedup", methods=["POST"])
def admin_dedup():
    """Remove cross-portal duplicate jobs, keeping highest-scoring copy."""
    require_admin()
    deleted = dedup_jobs()
    return jsonify({"ok": True, "deleted": deleted})


@app.route("/preferences", methods=["GET", "POST"])
def preferences():
    """
    Simplified settings page: job preferences + account + sign-out only.
    The other fields (email creds, telegram, linkedin, apollo, agent
    config) still exist in storage — they're managed via .env / Vercel
    project env now, not via UI — so the POST handler MERGES form values
    into the stored prefs instead of replacing the whole dict.
    """
    if request.method == "POST":
        uid = current_user_id()
        existing = load_preferences(uid) or DEFAULT_PREFS.copy()
        updated = dict(existing)

        # Job preferences (the only fields this UI exposes)
        if "job_titles" in request.form:
            updated["job_titles"] = [
                t.strip() for t in request.form.get("job_titles", "").split(",") if t.strip()
            ]
        if "locations" in request.form:
            updated["locations"] = [
                l.strip() for l in request.form.get("locations", "").split(",") if l.strip()
            ]
        if "transferable_skills" in request.form:
            updated["transferable_skills"] = [
                s.strip() for s in request.form.get("transferable_skills", "").split(",") if s.strip()
            ]
        if "industries" in request.form:
            updated["industries"] = [
                s.strip() for s in request.form.get("industries", "").split(",") if s.strip()
            ]
        work_modes = request.form.getlist("work_modes")
        if work_modes:
            updated["work_modes"] = work_modes
        save_preferences(updated, user_id=uid)
        cv = load_cv_data(uid) if uid else None
        if uid and cv:
            try:
                _rescore_all_jobs(cv, user_id=uid, preferences=updated)
            except Exception as e:
                logger.warning("Post-prefs rescore failed for user %s: %s", uid, e)
        flash("Job preferences saved.", "success")
        return redirect(url_for("preferences"))

    uid = current_user_id()
    prefs = load_preferences(uid) or DEFAULT_PREFS.copy()
    if not prefs.get("transferable_skills"):
        prefs["transferable_skills"] = DEFAULT_PREFS["transferable_skills"]
    # Apply env overrides so the form reflects what will actually be used
    prefs = apply_env_overrides(dict(prefs))
    return render_template("preferences.html", prefs=prefs)


@app.route("/api/account/email-settings", methods=["POST"])
def api_save_email_settings():
    """Save optional digest recipient override (Google login email is the default)."""
    uid = require_user_id()
    data = request.get_json(silent=True) or {}
    existing = load_preferences(uid) or DEFAULT_PREFS.copy()
    updated = dict(existing)
    if "email" in data:
        updated["email"] = (data["email"] or "").strip()
    try:
        save_preferences(updated, user_id=uid)
    except Exception as e:
        logger.error("Failed to save email settings for user %s: %s", uid, e)
        return jsonify({"ok": False, "error": "Failed to save"}), 500
    return jsonify({"ok": True})


@app.route("/api/preferences", methods=["POST"])
def api_save_preferences():
    """AJAX endpoint used by the preferences modal on the Jobs page."""
    uid = current_user_id()
    if not uid:
        return jsonify({"ok": False, "error": "Not authenticated"}), 401
    data = request.get_json(silent=True) or {}
    existing = load_preferences(uid) or DEFAULT_PREFS.copy()
    updated = dict(existing)
    if "job_titles" in data:
        updated["job_titles"] = [t.strip() for t in data["job_titles"] if t.strip()]
    if "locations" in data:
        updated["locations"] = [l.strip() for l in data["locations"] if l.strip()]
    if "work_modes" in data:
        updated["work_modes"] = data["work_modes"] or list({"on-site", "remote", "hybrid"})
    if "transferable_skills" in data:
        updated["transferable_skills"] = [s.strip() for s in data["transferable_skills"] if s.strip()]
    if "industries" in data:
        updated["industries"] = [s.strip() for s in data["industries"] if s.strip()]
    try:
        save_preferences(updated, user_id=uid)
    except Exception as e:
        logger.error("Failed to save preferences for user %s: %s", uid, e)
        return jsonify({"ok": False, "error": "Failed to save"}), 500
    cv = load_cv_data(uid) if uid else None
    if uid and cv:
        try:
            _rescore_all_jobs(cv, user_id=uid, preferences=updated)
        except Exception as e:
            logger.warning("Post-prefs rescore failed for user %s: %s", uid, e)
    return jsonify({"ok": True})


@app.route("/api/jobs/<job_id>/tailored-points")
def tailored_points(job_id):
    """Generate tailored resume bullet points for a specific job."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM job_listings WHERE job_id = ?", (job_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        return jsonify({"ok": False, "error": "Job not found"}), 404
    job = dict(row)
    config = load_config()
    uid = current_user_id()
    preferences = apply_env_overrides(load_preferences(uid) or DEFAULT_PREFS.copy())
    cv = get_user_cv_data(uid) if uid else None
    points = generate_tailored_points(job, preferences, config, cv_data=cv)
    return jsonify({"ok": True, "points": points})



@app.route("/api/jobs/import", methods=["POST"])
@csrf.exempt
def import_jobs():
    """Accept scraped jobs from external sources (e.g. GitHub Actions)."""
    # Auth: check header first, before parsing body
    import_secret = os.environ.get("IMPORT_SECRET", "")
    provided = request.headers.get("X-Import-Secret", "")
    if not provided:
        # Fall back to body only after header check fails (backward compat)
        provided = (request.get_json(silent=True, force=True) or {}).get("secret", "")
    if import_secret and provided != import_secret:
        return jsonify({"ok": False, "error": "Unauthorized"}), 401

    data = request.get_json(silent=True) or {}

    jobs = data.get("jobs")
    if not jobs or not isinstance(jobs, list):
        return jsonify({"ok": False, "error": "Missing or invalid 'jobs' array"}), 400

    # Per-user CV scoring is deferred to the read-side lazy scorer.
    for job in jobs:
        job["job_id"] = generate_job_id(
            job.get("portal", "unknown"),
            job.get("company", ""),
            job.get("role", ""),
            job.get("location", ""),
        )
    inserted, skipped = insert_jobs_bulk(jobs)
    logger.info("Import API: inserted=%d, skipped=%d (total submitted=%d)", inserted, skipped, len(jobs))

    # Telegram alerts for qualified jobs
    preferences = apply_env_overrides(load_preferences() or DEFAULT_PREFS.copy())
    tg_token = preferences.get("telegram_bot_token", "").strip()
    tg_chat = preferences.get("telegram_chat_id", "").strip()
    tg_min = int(preferences.get("telegram_min_score", 65))
    alert_count = 0
    if tg_token and tg_chat:
        for job in jobs:
            if job.get("relevance_score", 0) >= tg_min:
                try:
                    send_telegram_alert(job, tg_token, tg_chat)
                    alert_count += 1
                except Exception as e:
                    logger.warning("Telegram alert failed for %s: %s", job.get("job_id"), e)
        if alert_count > 0 or inserted > 0:
            try:
                send_telegram_batch_summary(len(jobs), alert_count, inserted, tg_token, tg_chat)
            except Exception as e:
                logger.warning("Telegram batch summary failed: %s", e)

    return jsonify({"ok": True, "inserted": inserted, "skipped": skipped, "alerts": alert_count})


@app.route("/api/portals/update", methods=["POST"])
def update_portals():
    """Enable or disable job portals. Persists to config.json."""
    require_admin()
    data = request.get_json(force=True) or {}
    enabled_portals = data.get("enabled", [])   # list of portal names to enable

    config_path = os.path.join(BASE_DIR, "config.json")
    try:
        with open(config_path) as f:
            config = json.load(f)
    except Exception:
        config = load_config()

    all_portal_names = list(config.get("portals", {}).keys())
    for name in all_portal_names:
        config["portals"][name]["enabled"] = (name in enabled_portals)

    # Atomic write — temp file then rename so a crash never corrupts config.json
    tmp_path = config_path + ".tmp"
    with open(tmp_path, "w") as f:
        json.dump(config, f, indent=2)
    os.replace(tmp_path, config_path)

    return jsonify({"ok": True, "enabled": enabled_portals})


@app.route("/api/scraper/start", methods=["POST"])
def start_scraper():
    require_admin()
    if _IS_VERCEL:
        return jsonify({"ok": False, "error": "Scraper is not available in cloud mode. Run the scraper locally with: python main.py"}), 503
    global scraper_status
    with scraper_lock:
        if scraper_status["running"]:
            return jsonify({"ok": False, "error": "Scraper is already running"}), 409
        scraper_status = {
            "running": True,
            "phase": "starting",
            "portal_progress": {},
            "done_portals": 0,
            "total_portals": 0,
            "total_jobs": 0,
            "qualified_jobs": 0,
            "inserted": 0,
            "skipped": 0,
            "digest_path": None,
            "error": None,
            "started_at": datetime.now().isoformat(),
            "finished_at": None,
        }
    _scraper_stop_event.clear()
    _is_scheduled_run = False
    t = threading.Thread(target=_run_scraper_pipeline, daemon=True)
    t.start()
    return jsonify({"ok": True})


@app.route("/api/scraper/quick-start", methods=["POST"])
def quick_scraper():
    """
    Vercel-safe quick scrape — runs only HiringCafe, Remotive, and Hacker News.
    No Selenium required. Completes in ~20s. Available on both local and Vercel.
    """
    uid = require_admin()
    from scrapers import VERCEL_SAFE_PORTALS

    config = load_config()
    preferences = apply_env_overrides(load_preferences(uid) or DEFAULT_PREFS.copy())
    job_titles = preferences.get("job_titles", ["Product Manager"])
    locations  = preferences.get("locations", ["India"])

    try:
        from scrapers import scrape_all_portals as _scrape
        all_jobs, portal_results = _scrape(
            job_titles, locations, config,
            allowed_portals=VERCEL_SAFE_PORTALS,
        )
    except Exception as e:
        logger.error("Quick scrape failed: %s", e)
        return jsonify({"ok": False, "error": str(e)}), 500

    if not all_jobs:
        counts = {p: r["count"] for p, r in portal_results.items()}
        return jsonify({"ok": True, "inserted": 0, "skipped": 0, "portals": counts,
                        "message": "No new jobs found."})

    # Per-user scoring is handled by the lazy scorer on next /api/jobs hit.
    for job in all_jobs:
        job["job_id"] = generate_job_id(
            job.get("portal", "unknown"),
            job.get("company", ""),
            job.get("role", ""),
            job.get("location", ""),
        )

    inserted, skipped = insert_jobs_bulk(all_jobs)
    logger.info("Quick scrape: inserted=%d skipped=%d portals=%s", inserted, skipped,
                list(portal_results.keys()))
    counts = {p: r["count"] for p, r in portal_results.items()}
    return jsonify({"ok": True, "inserted": inserted, "skipped": skipped, "portals": counts})


@app.route("/api/scraper/trigger-github", methods=["POST"])
def trigger_github_scrape():
    """
    Dispatch the GitHub Actions scrape workflow via the GitHub REST API.
    Requires GITHUB_TOKEN (PAT with workflow scope) and GITHUB_REPO env vars.
    """
    require_admin()
    import requests as _req

    token = os.environ.get("GITHUB_TOKEN", "")
    repo  = os.environ.get("GITHUB_REPO", "gmpro-cr/Job-Search-Agent")
    if not token:
        return jsonify({"ok": False,
                        "error": "GITHUB_TOKEN not set. Add a Personal Access Token "
                                 "with 'workflow' scope to your environment variables."}), 400

    url = f"https://api.github.com/repos/{repo}/actions/workflows/scrape.yml/dispatches"
    try:
        resp = _req.post(
            url,
            json={"ref": "main"},
            headers={
                "Authorization": f"Bearer {token}",
                "Accept": "application/vnd.github+json",
                "X-GitHub-Api-Version": "2022-11-28",
            },
            timeout=10,
        )
    except Exception as e:
        return jsonify({"ok": False, "error": f"GitHub API request failed: {e}"}), 500

    if resp.status_code == 204:
        runs_url = f"https://github.com/{repo}/actions/workflows/scrape.yml"
        return jsonify({"ok": True, "runs_url": runs_url,
                        "message": "Scrape workflow triggered. Results arrive in ~10 minutes."})
    else:
        return jsonify({"ok": False, "error": f"GitHub API returned {resp.status_code}: {resp.text[:200]}"}), 500


@app.route("/api/scraper/stop", methods=["POST"])
def stop_scraper():
    """Stop a manually triggered scraper run. Does NOT affect scheduled runs."""
    require_admin()
    global scraper_status
    with scraper_lock:
        if not scraper_status["running"]:
            return jsonify({"ok": False, "error": "Scraper is not running"}), 409
        if _is_scheduled_run:
            return jsonify({"ok": False, "error": "A scheduled run is in progress. Use Stop Scheduled Run to cancel it."}), 409
    _scraper_stop_event.set()
    with scraper_lock:
        scraper_status["phase"] = "stopping"
    return jsonify({"ok": True})


@app.route("/api/scraper/stop-scheduled", methods=["POST"])
def stop_scheduled_scraper():
    """Stop the currently running scheduled scraper run."""
    require_admin()
    global scraper_status
    with scraper_lock:
        if not scraper_status["running"]:
            return jsonify({"ok": False, "error": "No scraper run in progress"}), 409
        if not _is_scheduled_run:
            return jsonify({"ok": False, "error": "No scheduled run in progress. Use Stop to cancel a manual run."}), 409
    _scheduled_stop_event.set()
    with scraper_lock:
        scraper_status["phase"] = "stopping"
    return jsonify({"ok": True})


@app.route("/api/scraper/status")
def scraper_status_api():
    with scraper_lock:
        data = dict(scraper_status)
    data["is_scheduled_run"] = _is_scheduled_run
    return jsonify(data)


# ---------------------------------------------------------------------------
# Live Search API
# ---------------------------------------------------------------------------

@app.route("/api/search/start", methods=["POST"])
def start_live_search():
    require_admin()
    if _IS_VERCEL:
        return jsonify({"ok": False, "error": "Live search is not available in cloud mode. Run the scraper locally with: python main.py"}), 503
    global live_search_status
    data = request.get_json(silent=True) or {}
    query = data.get("query", "").strip()
    location = data.get("location", "").strip()

    with live_search_lock:
        if live_search_status["running"]:
            return jsonify({"ok": False, "error": "A search is already running"}), 409
        live_search_status = {
            "running": True,
            "phase": "starting",
            "portal_progress": {},
            "done_portals": 0,
            "total_portals": 0,
            "total_jobs": 0,
            "qualified_jobs": 0,
            "inserted": 0,
            "skipped": 0,
            "error": None,
            "started_at": datetime.now().isoformat(),
            "finished_at": None,
            "result_job_ids": [],
        }
    t = threading.Thread(target=_run_live_search, args=(query, location), daemon=True)
    t.start()
    return jsonify({"ok": True})


@app.route("/api/search/status")
def live_search_status_api():
    with live_search_lock:
        return jsonify(dict(live_search_status))


# ---------------------------------------------------------------------------
# Scheduler & Digests
# ---------------------------------------------------------------------------

@app.route("/api/scheduler/status")
def scheduler_status():
    """Return the current scheduler state and next run time."""
    if _scheduler and _scheduler.running:
        morning_job = _scheduler.get_job("morning_pipeline")
        evening_job = _scheduler.get_job("evening_pipeline")
        job = morning_job or evening_job
        if job:
            next_run = job.next_run_time
            return jsonify({
                "enabled": True,
                "next_run": next_run.isoformat() if next_run else None,
                "next_run_human": next_run.strftime("%B %d, %Y at %I:%M %p") if next_run else None,
            })
    return jsonify({"enabled": False, "next_run": None, "next_run_human": None})


def _hm_load_json() -> dict:
    return _load_hm().load_hm_contacts_dict()


def _hm_save_json(data: dict) -> None:
    _load_hm().save_hm_contacts_dict(data)


def _load_all_hm_contacts() -> list:
    return _load_hm().load_all_hm_contacts()


@app.route("/hiring-managers")
def hiring_managers():
    """Recruiter/TA contacts found via LinkedIn search, sourced from hr_sent_contacts.json."""
    contacts = [c for c in _load_all_hm_contacts() if c.get("linkedin_url")]
    preferences = apply_env_overrides(load_preferences(current_user_id()) or DEFAULT_PREFS.copy())
    user_pref_title = (preferences.get("job_titles") or ["a relevant role"])[0]
    return render_template(
        "hiring_managers.html",
        contacts=contacts,
        total=len(contacts),
        user_pref_title=user_pref_title,
    )


@app.route("/api/hiring-managers/search", methods=["POST"])
def api_hiring_managers_search():
    """Run a fresh hiring manager search and store results (does not send email)."""
    uid = require_user_id()
    prefs = apply_env_overrides(load_preferences(uid) or DEFAULT_PREFS.copy())
    cv_data_r = load_cv_data(uid) or {}
    role_keywords = [t.strip() for t in (prefs.get("job_titles") or []) if t.strip()]
    if not role_keywords:
        return jsonify({"ok": False, "error": "No job titles configured in preferences"}), 400
    skills = cv_data_r.get("skills") or []
    location = prefs.get("preferred_location") or "India"
    rid = f"user_{uid}"
    try:
        hm = _load_hm()
        sent = hm.load_hr_sent(rid)
        new_contacts = hm.get_new_hiring_managers(
            sent, role_keywords=role_keywords, skills=skills,
            location=location, target=10,
        )
        if new_contacts:
            existing = _hm_load_json()
            existing.setdefault(rid, [])
            for c in new_contacts:
                existing[rid].append({
                    "name": c["name"], "company": c["company"],
                    "their_role": c.get("their_role", ""),
                    "linkedin_url": c.get("linkedin_url", ""),
                    "date_sent": date.today().isoformat(),
                })
            _hm_save_json(existing)
        return jsonify({"ok": True, "found": len(new_contacts)})
    except Exception as e:
        logger.error("Hiring manager search failed: %s", e)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/digests")
def digests():
    uid = current_user_id()
    prefs = load_preferences(uid)
    from database import get_user_reminders
    job_alerts = get_user_reminders(uid) if uid else []
    user_email = dict(session).get('user', {}).get('email', '')
    return render_template("digests.html", prefs=prefs,
                           user_email=user_email, job_alerts=job_alerts)


@app.route("/api/digest/send-now", methods=["POST"])
def api_digest_send_now():
    """Send the user's top matched jobs by email right now."""
    uid = require_user_id()
    if _is_demo_user():
        return jsonify({"ok": False, "error": "Email sending is disabled in demo mode."}), 403

    # Sender credentials come from env vars only (app-level, not per-user).
    app_prefs = apply_env_overrides({})
    gmail_address = app_prefs.get("gmail_address", "").strip()
    gmail_app_password = app_prefs.get("gmail_app_password", "").strip()
    if not gmail_address or not gmail_app_password:
        return jsonify({"ok": False, "error": "Email sending is not configured on this server. Contact the administrator."}), 503

    # If the user has saved routines, "Send now" runs THOSE (one email per
    # routine, each with its own keywords/location/score), matching the
    # automation — forced (ignore last_sent) so it always sends current matches.
    from database import get_user_reminders
    active_routines = [
        r for r in (get_user_reminders(uid) or [])
        if r.get("enabled", True) and (r.get("email") or "").strip()
        and (r.get("keyword") or "").strip()
    ]
    if active_routines:
        from routine_runner import run_routines
        rr = run_routines(user_id=uid)
        if rr["emails_sent"]:
            n = rr["emails_sent"]
            return jsonify({"ok": True,
                            "message": f"Sent {n} routine email{'' if n == 1 else 's'} "
                                       f"({rr['jobs_sent']} jobs)."})
        return jsonify({"ok": False,
                        "error": "No fresh matches for your routines right now."}), 400

    # No routines: fall back to the default top-matches digest.
    # Load once; apply_env_overrides injects sender creds from env, not user storage.
    prefs = apply_env_overrides(load_preferences(uid) or DEFAULT_PREFS.copy())
    google_email = (session.get("user") or {}).get("email", "")
    recipient = (prefs.get("email") or google_email or "").strip()
    if not recipient:
        return jsonify({"ok": False, "error": "Could not determine your email address. Please sign out and sign in again."}), 400

    top_n = max(1, min(50, int(prefs.get("top_jobs_per_digest", 10))))
    min_score = max(0, int(prefs.get("min_score", 50)))
    # Pull a wider candidate window so we can drop jobs already emailed to this
    # user in the last 7 days and still fill up to top_n with fresh ones.
    candidate_limit = min(200, max(top_n * 4, top_n))
    conn = get_connection()
    cursor = conn.cursor()
    if uid:
        cursor.execute(
            """
            SELECT j.*, s.cv_score AS user_cv_score
            FROM job_listings j
            JOIN user_job_state s ON s.job_id = j.job_id AND s.user_id = ?
            WHERE s.cv_score >= ? AND (COALESCE(s.hidden, 0) = 0)
            ORDER BY s.cv_score DESC
            LIMIT ?
            """,
            (uid, min_score, candidate_limit),
        )
    else:
        cursor.execute(
            "SELECT * FROM job_listings WHERE cv_score >= ? AND (hidden = 0 OR hidden IS NULL) ORDER BY cv_score DESC LIMIT ?",
            (min_score, candidate_limit),
        )
    candidates = [dict(r) for r in cursor.fetchall()]
    conn.close()

    # Prefer jobs not emailed to this user recently, but backfill so a manual
    # send is never empty when few jobs qualify.
    jobs = select_digest_jobs(candidates, top_n, days=7, user_id=uid)
    if not jobs:
        return jsonify({"ok": False, "error": "No matching jobs found to send."}), 400

    from email_notifier import send_job_email
    send_prefs = dict(prefs)
    send_prefs["gmail_address"] = gmail_address
    send_prefs["gmail_app_password"] = gmail_app_password
    ok, err = send_job_email(recipient, jobs, send_prefs)
    if ok:
        try:
            mark_sent_in_digest(
                [j.get("job_id") for j in jobs if j.get("job_id")], user_id=uid,
            )
        except Exception as e:
            logger.warning("mark_sent_in_digest failed for user %s: %s", uid, e)
        return jsonify({"ok": True, "message": f"Sent {len(jobs)} jobs to {recipient}"})
    return jsonify({"ok": False, "error": err or "Email send failed."}), 500


@app.route("/api/digest/settings", methods=["POST"])
def api_digest_settings():
    """Save digest-specific preferences (email, top_n, min_score)."""
    uid = current_user_id()
    if not uid:
        return jsonify({"ok": False, "error": "Not authenticated"}), 401
    data = request.get_json(silent=True) or {}
    existing = load_preferences(uid) or DEFAULT_PREFS.copy()
    updated = dict(existing)
    if "email" in data:
        updated["email"] = (data["email"] or "").strip()
    if "top_jobs_per_digest" in data:
        try:
            updated["top_jobs_per_digest"] = max(1, min(50, int(data["top_jobs_per_digest"])))
        except (ValueError, TypeError):
            pass
    if "min_score" in data:
        try:
            updated["min_score"] = max(0, min(100, int(data["min_score"])))
        except (ValueError, TypeError):
            pass
    try:
        save_preferences(updated, user_id=uid)
    except Exception as e:
        logger.error("Failed to save digest settings for user %s: %s", uid, e)
        return jsonify({"ok": False, "error": "Failed to save"}), 500
    return jsonify({"ok": True})


@app.route("/api/digest/deactivate", methods=["POST"])
def api_digest_deactivate():
    """Clear the digest email so no more digests are sent."""
    uid = current_user_id()
    if not uid:
        return jsonify({"ok": False, "error": "Not authenticated"}), 401
    existing = load_preferences(uid) or DEFAULT_PREFS.copy()
    updated = dict(existing)
    updated["email"] = ""
    try:
        save_preferences(updated, user_id=uid)
    except Exception as e:
        logger.error("Failed to deactivate digest for user %s: %s", uid, e)
        return jsonify({"ok": False, "error": "Failed to save"}), 500
    return jsonify({"ok": True})



# ---------------------------------------------------------------------------
# CV Management Routes
# ---------------------------------------------------------------------------

@app.route("/cv", methods=["GET", "POST"])
def cv_page():
    if request.method == "POST":
        if "cv_file" not in request.files:
            flash("No file selected.", "error")
            return redirect(url_for("cv_page"))
        f = request.files["cv_file"]
        if not f.filename:
            flash("No file selected.", "error")
            return redirect(url_for("cv_page"))
        filename = f.filename or ""
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        text = ""
        if ext == "pdf":
            try:
                import pdfplumber, io as _io
                with pdfplumber.open(_io.BytesIO(f.read())) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            except Exception as e:
                flash(f"PDF parsing failed: {e}", "error")
                return redirect(url_for("cv_page"))
        elif ext == "docx":
            try:
                import docx as _docx, io as _io
                doc = _docx.Document(_io.BytesIO(f.read()))
                text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                flash(f"DOCX parsing failed: {e}", "error")
                return redirect(url_for("cv_page"))
        else:
            text = f.read().decode("utf-8", errors="ignore")
        _uid = require_user_id()
        cv_data = parse_cv_text(text)
        cv_data["filename"] = filename
        save_cv_data(cv_data, user_id=_uid)
        # Auto-populate any empty preference fields (titles / locations /
        # skills) from the parsed CV so first-time users land on a
        # working dashboard without filling forms.
        merged_prefs = _autofill_prefs_from_cv(cv_data, user_id=_uid)
        _schedule_rescore(cv_data, merged_prefs, _uid)
        flash(f"CV uploaded — {len(cv_data['skills'])} skills detected. Scoring jobs in the background…", "success")
        _next = request.form.get("next", "")
        from urllib.parse import urlparse as _urlparse
        if not _next or _urlparse(_next).netloc:
            _next = url_for("cv_page")
        return redirect(_next)

    uid = require_user_id()
    cv_data = load_cv_data(uid)

    cv_filename = cv_data.get("filename") if cv_data else None
    cv_uploaded_at = cv_data.get("uploaded_at") if cv_data else None
    skills = cv_data.get("skills", []) if cv_data else []
    cv_skills_count = len(skills)
    skill_lower = [s.lower() for s in skills]

    conn = get_connection()
    cur = conn.cursor()

    # Total job count for demand %
    cur.execute("SELECT COUNT(*) as n FROM job_listings")
    total_job_count = max(cur.fetchone()["n"] or 1, 1)

    # Single JD scan
    cur.execute("SELECT job_description FROM job_listings LIMIT 3000")
    all_jds = " ".join((r["job_description"] or "") for r in cur.fetchall()).lower()

    # Skill demand: CV skills sorted by frequency in job market
    skill_demand = []
    for sk in skills[:14]:
        freq = all_jds.count(sk.lower())
        pct = round(freq / total_job_count * 100)
        skill_demand.append({"name": sk, "frequency": freq, "pct": min(100, pct)})
    skill_demand.sort(key=lambda x: -x["frequency"])

    # Skill gaps: broad domain-relevant list, not in CV
    _common = [
        "Machine Learning", "Generative AI", "LLMs", "Product Strategy", "Data Analysis",
        "SQL", "Python", "Stakeholder Management", "Agile", "Scrum", "AWS", "Docker",
        "System Design", "API Design", "User Research", "A/B Testing", "Product Analytics",
        "Roadmapping", "Go-to-Market", "Growth", "NLP", "React", "Kubernetes",
        "TypeScript", "Figma", "OKRs", "PRD Writing", "Leadership", "B2B SaaS",
    ]
    skill_gaps = []
    for sk in _common:
        if sk.lower() not in skill_lower:
            freq = all_jds.count(sk.lower())
            pct = round(freq / total_job_count * 100)
            if freq > 0:
                skill_gaps.append({"skill": sk, "demand_count": freq, "pct": min(100, pct)})
    skill_gaps.sort(key=lambda x: -x["demand_count"])
    skill_gaps = skill_gaps[:8]

    # Score distribution across all user-scored jobs
    score_dist_data = [0, 0, 0, 0, 0]  # 0-20, 21-40, 41-60, 61-80, 81-100
    total_scored = 0
    avg_match_score = 0
    if uid:
        cur.execute(
            "SELECT cv_score FROM user_job_state WHERE user_id = ? AND cv_score > 0",
            (uid,),
        )
        _scores = [r["cv_score"] for r in cur.fetchall()]
        total_scored = len(_scores)
        if _scores:
            avg_match_score = round(sum(_scores) / len(_scores))
            for sc in _scores:
                if sc <= 20:   score_dist_data[0] += 1
                elif sc <= 40: score_dist_data[1] += 1
                elif sc <= 60: score_dist_data[2] += 1
                elif sc <= 80: score_dist_data[3] += 1
                else:          score_dist_data[4] += 1

    # Portal breakdown (avg score + count)
    portal_breakdown = []
    if uid:
        cur.execute(
            """
            SELECT j.portal, COUNT(*) as cnt, AVG(s.cv_score) as avg_score
            FROM user_job_state s JOIN job_listings j ON j.job_id = s.job_id
            WHERE s.user_id = ? AND s.cv_score > 0
              AND j.portal IS NOT NULL AND j.portal != ''
            GROUP BY j.portal ORDER BY avg_score DESC LIMIT 6
            """,
            (uid,),
        )
        portal_breakdown = [
            {"portal": r["portal"], "count": r["cnt"],
             "avg_score": round(float(r["avg_score"] or 0))}
            for r in cur.fetchall()
        ]

    # Top matched job titles (cv_score >= 65)
    top_titles = []
    if uid:
        cur.execute(
            """
            SELECT j.role, COUNT(*) as cnt FROM user_job_state s
            JOIN job_listings j ON j.job_id = s.job_id
            WHERE s.user_id = ? AND s.cv_score >= 65
              AND j.role IS NOT NULL AND j.role != ''
            GROUP BY j.role ORDER BY cnt DESC LIMIT 8
            """,
            (uid,),
        )
        top_titles = [{"title": r["role"], "count": r["cnt"]} for r in cur.fetchall()]

    conn.close()

    # Skills tag cloud: (name, demand_pct) sorted by demand
    tag_cloud = sorted(
        [(sk, round(all_jds.count(sk.lower()) / total_job_count * 100)) for sk in skills[:24]],
        key=lambda x: -x[1],
    )

    return render_template("cv.html",
        cv_data=cv_data,
        cv_filename=cv_filename,
        cv_uploaded=cv_data is not None,
        cv_uploaded_at=cv_uploaded_at,
        cv_skills_count=cv_skills_count,
        skill_demand=skill_demand,
        skill_gaps=skill_gaps,
        score_dist_data=score_dist_data,
        total_scored=total_scored,
        avg_match_score=avg_match_score,
        portal_breakdown=portal_breakdown,
        top_titles=top_titles,
        tag_cloud=tag_cloud,
    )


@app.route("/api/cv/upload", methods=["POST"])
def upload_cv():
    """Accept a CV file upload, parse it, and store it for the current user."""
    uid = require_user_id()
    if "cv_file" not in request.files:
        return jsonify({"ok": False, "error": "No file provided"}), 400

    f = request.files["cv_file"]
    filename = f.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    text = ""
    if ext == "pdf":
        try:
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(f.read())) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as e:
            return jsonify({"ok": False, "error": f"PDF parsing failed: {e}"}), 400
    elif ext == "docx":
        try:
            import docx, io
            doc = docx.Document(io.BytesIO(f.read()))
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            return jsonify({"ok": False, "error": f"DOCX parsing failed: {e}"}), 400
    elif ext in ("txt", ""):
        text = f.read().decode("utf-8", errors="ignore")
    else:
        return jsonify({"ok": False, "error": f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT."}), 400

    if not text.strip():
        return jsonify({"ok": False, "error": "Could not extract text from the file"}), 400

    cv_data = parse_cv_text(text)
    cv_data["filename"] = filename
    save_cv_data(cv_data, user_id=uid)
    logger.info("CV uploaded by user %s: %d skills detected", uid, len(cv_data["skills"]))

    # Auto-fill preferences from the CV (only empty fields), then schedule a
    # rescore. On Vercel the rescore is deferred to the lazy path on the
    # next /api/jobs hit — synchronous full-table rescore would burn the
    # function timeout for a large job_listings table.
    merged_prefs = _autofill_prefs_from_cv(cv_data, user_id=uid)
    rescored = _schedule_rescore(cv_data, merged_prefs, uid)

    return jsonify({
        "ok": True,
        "skills_count": len(cv_data["skills"]),
        "skills": cv_data["skills"],
        "suggested_job_titles": cv_data.get("suggested_job_titles", []),
        "suggested_locations":  cv_data.get("suggested_locations", []),
        "rescored": rescored,
    })


def _rescore_in_background(cv_data, preferences, user_id):
    try:
        _rescore_all_jobs(cv_data, preferences=preferences, user_id=user_id)
    except Exception as e:
        logger.error("Background rescore failed for user %s: %s", user_id, e)


# Cap for inline rescores on Vercel — large job tables would otherwise
# exceed the function timeout. Remaining jobs get picked up lazily by
# _score_unscored_for_user on the next /api/jobs hit.
_VERCEL_INLINE_RESCORE_CAP = 500


def _schedule_rescore(cv_data, preferences, user_id):
    """Rescore after a CV/preferences change.

    - Local: fire-and-forget thread (existing behaviour).
    - Vercel: background threads get killed when the function returns,
      so we run an inline batch (capped) and let the lazy
      _score_unscored_for_user path fill in the rest on subsequent
      requests.

    Returns the number of jobs scored synchronously (always 0 on local
    since the thread runs after we return).
    """
    if not user_id:
        return 0
    if _IS_VERCEL:
        try:
            return _rescore_all_jobs(
                cv_data,
                preferences=preferences,
                user_id=user_id,
                limit=_VERCEL_INLINE_RESCORE_CAP,
            )
        except Exception as e:
            logger.error("Inline rescore failed for user %s: %s", user_id, e)
            return 0
    threading.Thread(
        target=_rescore_in_background,
        args=(cv_data, preferences, user_id),
        daemon=True,
    ).start()
    return 0


def _rescore_all_jobs(cv_data, user_id: int = None, preferences: dict = None,
                      limit: int = None):
    """Score jobs in the DB against cv_data + preferences and persist
    into user_job_state for the given user.

    Requires user_id in a Flask request context (the legacy global path
    only exists for CLI / scraper use).
    """
    if user_id is None:
        user_id = current_user_id()
    # Refuse to silently fall through to the legacy global column when we
    # have a Flask request context — that would clobber the owner's data.
    if user_id is None:
        try:
            from flask import has_request_context as _hrc
            if _hrc():
                raise RuntimeError(
                    "_rescore_all_jobs called from a Flask request with no user_id"
                )
        except ImportError:
            pass
    if preferences is None and user_id:
        preferences = load_preferences(user_id) or {}
    conn = get_connection()
    cursor = conn.cursor()
    if limit:
        cursor.execute(
            "SELECT job_id, role, company, location, job_description, remote_status "
            "FROM job_listings ORDER BY date_found DESC LIMIT ?",
            (int(limit),),
        )
    else:
        cursor.execute(
            "SELECT job_id, role, company, location, job_description, remote_status "
            "FROM job_listings"
        )
    jobs = [dict(r) for r in cursor.fetchall()]
    conn.close()

    if not jobs:
        return 0

    if user_id:
        scores = {job["job_id"]: cv_score(job, cv_data, preferences) for job in jobs}
        bulk_set_user_cv_scores(user_id, scores)
        logger.info("Re-scored %d jobs against CV+prefs for user %s", len(jobs), user_id)
        return len(jobs)

    conn = get_connection()
    cursor = conn.cursor()
    for job in jobs:
        score = cv_score(job, cv_data)
        cursor.execute(
            "UPDATE job_listings SET cv_score = ? WHERE job_id = ?",
            (score, job["job_id"]),
        )
    conn.commit()
    conn.close()
    logger.info("Re-scored %d jobs against CV (legacy global)", len(jobs))
    return len(jobs)


def _autofill_prefs_from_cv(cv_data: dict, user_id: int = None) -> dict:
    """
    After CV parse, fill any empty preference fields with what we
    inferred from the CV. Never overwrites a field the user has
    already populated — purely additive.
    Returns the merged preferences dict that was saved.
    """
    if user_id is None:
        user_id = current_user_id()
    if not user_id or not cv_data:
        return {}
    existing = load_preferences(user_id) or {}
    merged = dict(existing)

    suggestions = {
        "job_titles":          cv_data.get("suggested_job_titles") or [],
        "locations":           cv_data.get("suggested_locations") or [],
        "transferable_skills": cv_data.get("skills") or [],
    }
    changed = False
    for key, suggestion in suggestions.items():
        existing_list = merged.get(key) or []
        if not existing_list and suggestion:
            # Keep the order, take up to 6 to avoid bloating the form.
            merged[key] = suggestion[:6]
            changed = True

    if changed:
        save_preferences(merged, user_id=user_id)
        logger.info("Auto-filled prefs from CV for user %s: %s",
                    user_id,
                    {k: v for k, v in merged.items()
                     if k in ("job_titles", "locations", "transferable_skills")})
    return merged


def _score_unscored_for_user(user_id: int, limit: int = 200) -> int:
    """
    Lazy per-user scorer. On each /api/jobs hit (or wherever called),
    find jobs that don't yet have a user_job_state row for this user,
    score them against the user's CV + preferences, and bulk-write
    the resulting cv_scores.

    No-op if the user has no CV uploaded — without a CV every score
    would be 0 and we'd just be writing zero rows.

    Returns: number of jobs scored.
    """
    if not user_id:
        return 0
    cv_data = load_cv_data(user_id)
    if not cv_data:
        return 0  # nothing to score against
    from database import get_unscored_jobs_for_user
    jobs = get_unscored_jobs_for_user(user_id, limit=limit)
    if not jobs:
        return 0
    preferences = load_preferences(user_id) or {}
    scores = {j["job_id"]: cv_score(j, cv_data, preferences) for j in jobs}
    bulk_set_user_cv_scores(user_id, scores)
    logger.info("Lazy-scored %d new jobs for user %s", len(jobs), user_id)
    return len(jobs)


@app.route("/api/cv/skills")
def api_cv_skills():
    """Return parsed CV skills and raw text."""
    uid = require_user_id()
    cv_data = load_cv_data(uid)
    if not cv_data:
        return jsonify({"ok": False, "error": "CV not uploaded yet"}), 400
    return jsonify({"ok": True, "skills": cv_data.get("skills", []), "raw_text": cv_data.get("raw_text", "")})

@app.route("/api/cv/top_jobs")
def api_top_jobs():
    """Return top 10 jobs matched to the uploaded CV, sorted by cv_score descending."""
    uid = require_user_id()
    cv_data = load_cv_data(uid)
    if not cv_data:
        return jsonify({"ok": False, "error": "CV not uploaded yet"}), 400
    conn = get_connection()
    cursor = conn.cursor()
    if uid:
        cursor.execute(
            """
            SELECT j.job_id, j.role, j.company, j.apply_url,
                   COALESCE(s.cv_score, 0) AS cv_score
            FROM job_listings j
            JOIN user_job_state s ON s.job_id = j.job_id AND s.user_id = ?
            WHERE s.cv_score IS NOT NULL AND s.cv_score > 0
            ORDER BY s.cv_score DESC
            LIMIT 10
            """,
            (uid,),
        )
    else:
        cursor.execute(
            "SELECT job_id, role, company, apply_url, cv_score FROM job_listings "
            "WHERE cv_score IS NOT NULL ORDER BY cv_score DESC LIMIT 10"
        )
    rows = [dict(r) for r in cursor.fetchall()]
    conn.close()
    return jsonify({"ok": True, "jobs": rows})

@app.route("/api/cv/rescore", methods=["POST"])
def rescore_jobs():
    """Re-score all jobs in the DB against the uploaded CV."""
    uid = require_user_id()
    cv_data = load_cv_data(uid)
    if not cv_data:
        return jsonify({"ok": False, "error": "No CV uploaded yet"}), 400
    preferences = load_preferences(uid) or {}
    rescored = _schedule_rescore(cv_data, preferences, uid)
    return jsonify({"ok": True, "updated": rescored})


@app.route("/api/cv/skills-gap")
def cv_skills_gap():
    """Return skill frequency across target-role jobs vs CV skills."""
    uid = require_user_id()
    cv_data = load_cv_data(uid)
    # If no CV data, redirect to upload page
    if not cv_data:
        return redirect(url_for('cv_page'))
    preferences = load_preferences(uid) or DEFAULT_PREFS.copy()
    job_titles = preferences.get("job_titles", [])

    from database import get_skill_frequency
    skill_freq = get_skill_frequency(job_titles)

    cv_skills_lower = {s.lower() for s in (cv_data or {}).get("skills", [])}
    result = []
    for item in skill_freq:
        result.append({
            "skill": item["skill"],
            "count": item["count"],
            "pct": item["pct"],
            "in_cv": item["skill"].lower() in cv_skills_lower,
        })
    return jsonify({"ok": True, "skills": result, "cv_uploaded": cv_data is not None})


@app.route("/api/cv/keyword-heatmap")
def cv_keyword_heatmap():
    """Return top keywords across target-role jobs."""
    uid = require_user_id()
    preferences = load_preferences(uid) or DEFAULT_PREFS.copy()
    job_titles = preferences.get("job_titles", [])

    from database import get_keyword_frequency
    keywords = get_keyword_frequency(job_titles)
    return jsonify({"ok": True, "keywords": keywords})


@app.route("/api/cv/ats-score")
def cv_ats_score():
    """ATS-readiness score for the user's CV (heuristic; see analyzer.compute_ats_score)."""
    uid = require_user_id()
    cv_data = load_cv_data(uid)
    if not cv_data:
        return jsonify({"ok": True, "score": 0, "band": "No CV", "breakdown": [],
                        "suggestions": ["Upload your CV to get an ATS readiness score."]})
    preferences = load_preferences(uid) or DEFAULT_PREFS.copy()
    market_skills = []
    try:
        from database import get_skill_frequency
        market_skills = get_skill_frequency(preferences.get("job_titles") or [], limit=20)
    except Exception as e:
        logger.warning("ATS market-skill lookup failed for user %s: %s", uid, e)
    from analyzer import compute_ats_score
    result = compute_ats_score(cv_data, preferences, market_skills)
    result["ok"] = True
    return jsonify(result)


@app.route("/api/cv/profile-score")
def cv_profile_score():
    """Return a simple profile completeness score 0-100."""
    uid = require_user_id()
    cv_data = load_cv_data(uid)
    preferences = load_preferences(uid) or DEFAULT_PREFS.copy()

    score = 0
    breakdown = []

    if cv_data:
        score += 30
        breakdown.append({"label": "CV uploaded", "points": 30, "done": True})
    else:
        breakdown.append({"label": "Upload your CV", "points": 30, "done": False})

    skills = (cv_data or {}).get("skills", [])
    if len(skills) >= 5:
        score += 20
        breakdown.append({"label": f"{len(skills)} skills detected", "points": 20, "done": True})
    else:
        breakdown.append({"label": "CV needs more skills (aim for 5+)", "points": 20, "done": False})

    has_prefs = bool(preferences.get("job_titles") and preferences.get("locations"))
    if has_prefs:
        score += 20
        breakdown.append({"label": "Preferences configured", "points": 20, "done": True})
    else:
        breakdown.append({"label": "Set job titles & locations in Preferences", "points": 20, "done": False})

    has_salary = bool(preferences.get("salary_min") or preferences.get("salary_expectation"))
    if has_salary:
        score += 15
        breakdown.append({"label": "Salary expectation set", "points": 15, "done": True})
    else:
        breakdown.append({"label": "Add salary expectation in Preferences", "points": 15, "done": False})

    has_gmail = bool(preferences.get("gmail_address") and preferences.get("gmail_app_password"))
    if has_gmail:
        score += 15
        breakdown.append({"label": "Email alerts configured", "points": 15, "done": True})
    else:
        breakdown.append({"label": "Configure Gmail in Preferences for email alerts", "points": 15, "done": False})

    return jsonify({"ok": True, "score": score, "breakdown": breakdown})


# ── Outreach agent routes ────────────────────────────────────────────────────

@app.route("/api/outreach/<token>/save", methods=["POST"])
@csrf.exempt
def save_outreach_draft(token):
    """Save edited email/LinkedIn draft text for an outreach item."""
    from database import update_outreach_draft, get_outreach_by_token
    if get_outreach_by_token(token) is None:
        return jsonify({"ok": False, "error": "Not found"}), 404
    data = request.get_json(silent=True) or {}
    email_draft = data.get("email_draft")
    linkedin_draft = data.get("linkedin_draft")
    if email_draft is None and linkedin_draft is None:
        return jsonify({"ok": False, "error": "No draft field provided"}), 400
    updated = update_outreach_draft(token, email_draft=email_draft, linkedin_draft=linkedin_draft)
    return jsonify({"ok": updated})


@app.route("/api/approve/<token>", methods=["GET", "POST"])
@csrf.exempt
def approve_outreach(token):
    """Approve and send a cold email for the given token.

    GET renders a confirmation page (so link prefetchers, Slack/Twitter
    unfurl bots, anti-virus crawlers etc. cannot accidentally trigger a
    send). POST from that page performs the actual action.
    """
    from database import get_outreach_by_token, claim_outreach_token

    item = get_outreach_by_token(token)
    if item is None:
        return render_template("approve_result.html",
                               success=False, item=None,
                               message="This approval link is invalid."), 200

    if request.method == "GET":
        # Just show what would be sent; do nothing.
        return render_template("approve_confirm.html", item=item, token=token)

    # POST → claim + send.
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    claimed = claim_outreach_token(token)
    if claimed is None:
        return render_template("approve_result.html",
                               success=False, item=item,
                               message="This approval link has already been used or expired."), 200

    prefs = apply_env_overrides(load_preferences() or DEFAULT_PREFS.copy())
    gmail_address = prefs.get("gmail_address", "")
    gmail_password = prefs.get("gmail_app_password", "")

    if not gmail_address or not gmail_password:
        return "Gmail not configured. Please set it up in Preferences.", 500

    recipient_email = claimed.get("hm_email", "")
    if not recipient_email:
        return "No hiring manager email on file for this job.", 400

    try:
        apply_url = (claimed.get("apply_url") or "").strip()
        email_body = claimed["email_draft"]
        if apply_url:
            email_body += f"\n\nJob posting: {apply_url}"

        msg = MIMEMultipart("alternative")
        msg["Subject"] = f"Regarding the {claimed['role']} role at {claimed['company']}"
        msg["From"] = gmail_address
        msg["To"] = recipient_email
        msg.attach(MIMEText(email_body, "plain"))

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(gmail_address, gmail_password)
            smtp.sendmail(gmail_address, recipient_email, msg.as_string())

        # Mark as applied for the owner of the outreach draft.
        owner_uid = claimed.get("user_id")
        if owner_uid:
            update_applied_status_user(int(owner_uid), claimed["job_id"], 1)
        else:
            update_applied_status(claimed["job_id"], 1)
        return render_template("approve_result.html",
                               success=True, item=claimed,
                               message=f"Email sent to {recipient_email}!")
    except Exception as e:
        logger.error("approve_outreach send failed: %s", e)
        return f"Failed to send email: {e}", 500


@app.route("/api/skip/<token>", methods=["GET", "POST"])
@csrf.exempt
def skip_outreach(token):
    """Mark an outreach draft as skipped. GET confirms; POST acts."""
    from database import get_outreach_by_token, update_outreach_status

    item = get_outreach_by_token(token)
    if not item:
        return render_template("approve_result.html",
                               success=False, item=None,
                               message="Invalid or expired link."), 404

    if request.method == "GET":
        return render_template("approve_confirm.html",
                               item=item, token=token, skip=True)

    update_outreach_status(token, "skipped")
    return render_template("approve_result.html",
                           success=False, item=item,
                           message="Skipped. This job won't appear again.")


def _run_agent_background():
    """Run the AI agent pipeline in a background thread."""
    global agent_status
    try:
        import json as _json
        from agent.graph import run_agent_pipeline
        prefs = load_preferences() or DEFAULT_PREFS.copy()
        with open(os.path.join(BASE_DIR, "config.json")) as f:
            _config = _json.load(f)
        result = run_agent_pipeline(prefs, _config)
        with agent_lock:
            agent_status["queued"] = result.get("queued_count", 0)
            agent_status["error"] = result.get("errors", [None])[0] if result.get("errors") else None
    except Exception as e:
        logger.error("Agent background run error: %s", e)
        with agent_lock:
            agent_status["error"] = str(e)
    finally:
        with agent_lock:
            agent_status["running"] = False
            agent_status["finished_at"] = datetime.now().isoformat()


@app.route("/api/agent/run", methods=["POST"])
def run_agent_now():
    """Manually trigger the AI agent pipeline."""
    require_admin()
    global agent_status
    with agent_lock:
        if agent_status["running"]:
            return jsonify({"ok": False, "error": "Agent is already running"}), 409
        agent_status = {"running": True, "queued": 0, "error": None, "finished_at": None}
    t = threading.Thread(target=_run_agent_background, daemon=True)
    t.start()
    return jsonify({"ok": True})


@app.route("/api/agent/status")
def agent_run_status():
    """Return current agent run state."""
    with agent_lock:
        return jsonify(dict(agent_status))


@app.route("/api/jobs/<job_id>/gap-analysis")
def gap_analysis(job_id):
    """Return gap analysis for a specific job against the uploaded CV."""
    uid = require_user_id()
    cv_data = load_cv_data(uid)

    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM job_listings WHERE job_id = ?", (job_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        return jsonify({"ok": False, "error": "Job not found"}), 404

    job = dict(row)
    result = compute_gap_analysis(job, cv_data)
    return jsonify({"ok": True, **result})


@app.route("/api/jobs/<job_id>/score-breakdown")
def score_breakdown_api(job_id):
    """Return per-component score breakdown for a specific job."""
    uid = require_user_id()

    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM job_listings WHERE job_id = ?", (job_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "not found"}), 404

    job = dict(row)
    raw = (job.get("score_breakdown") or "")
    try:
        stored = json.loads(raw) if raw else {}
    except Exception:
        stored = {}

    # If we have a stored breakdown with real values, return it directly.
    if stored and any(v for v in stored.values() if isinstance(v, (int, float)) and v != 0):
        conn.close()
        return jsonify(stored)

    # No meaningful stored breakdown — compute from gap analysis so the
    # panel always shows something useful.
    cv_data = load_cv_data(uid)
    gap = compute_gap_analysis(job, cv_data)

    # Pull the stored cv_score for this user from user_job_state.
    cursor.execute(
        "SELECT cv_score FROM user_job_state WHERE user_id = ? AND job_id = ?",
        (uid, job_id),
    )
    ujs = cursor.fetchone()
    conn.close()
    stored_cv = (ujs["cv_score"] if ujs else 0) or 0

    return jsonify({
        "computed": True,
        "cv_score": stored_cv or gap.get("cv_score", 0),
        "skill_match": gap.get("cv_score", 0),
        "matched_skills": gap.get("matched_skills", []),
        "missing_skills": gap.get("missing_skills", []),
        "action_steps": gap.get("action_steps", []),
    })


def _extract_cv_text(file_storage) -> str:
    """
    Extract plain text from an uploaded CV file (PDF, DOCX, or TXT).
    Returns extracted text string, or raises ValueError with a user-friendly message.
    """
    filename = (file_storage.filename or "").lower()
    ext = filename.rsplit(".", 1)[-1] if "." in filename else ""
    data = file_storage.read()

    if ext == "pdf":
        try:
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as e:
            raise ValueError(f"PDF parsing failed: {e}")
    elif ext == "docx":
        try:
            import docx, io
            doc = docx.Document(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise ValueError(f"DOCX parsing failed: {e}")
    elif ext == "txt":
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type '.{ext}'. Use PDF, DOCX, or TXT.")

    if not text.strip():
        raise ValueError("Could not extract any text from the CV file.")
    return text


# ---------------------------------------------------------------------------
# Reminders
# ---------------------------------------------------------------------------

@app.route("/api/reminders", methods=["POST"])
def create_reminder():
    """Create a new saved routine for the current user."""
    uid = require_user_id()
    if isinstance(uid, tuple):
        return uid
    from database import get_user_reminders, save_user_reminders, uuid_safe
    data = request.get_json(silent=True) or {}
    keyword = str(data.get("keyword", "")).strip()[:500]
    if not keyword:
        return jsonify({"ok": False, "error": "Add at least one role or keyword."}), 400
    google_email = (session.get("user") or {}).get("email", "")
    email = (str(data.get("email", "")).strip() or google_email)[:200]
    if not email:
        return jsonify({"ok": False, "error": "Add a recipient email."}), 400
    new = {
        "id": uuid_safe(),
        "name": (str(data.get("name", "")).strip() or keyword)[:100],
        "email": email,
        "keyword": keyword,
        "location": str(data.get("location", "")).strip()[:300],
        "min_score": int(data.get("min_score", 50)),
        "max_jobs": int(data.get("max_jobs", 10)),
        "max_age_days": int(data.get("max_age_days", 0) or 0),
        "enabled": True,
        "last_sent": None,
    }
    reminders = get_user_reminders(uid)
    reminders.append(new)
    save_user_reminders(uid, reminders)
    return jsonify({"ok": True, "id": new["id"]})


@app.route("/api/reminders/<reminder_id>", methods=["PUT"])
def update_reminder(reminder_id):
    uid = require_user_id()
    if isinstance(uid, tuple):
        return uid
    from database import get_user_reminders, save_user_reminders
    data = request.get_json(silent=True) or {}
    reminders = get_user_reminders(uid)
    for r in reminders:
        if str(r.get("id")) == reminder_id:
            r["name"]      = str(data.get("name", r.get("name", "")))[:100]
            r["email"]     = str(data.get("email", r.get("email", "")))[:200]
            r["keyword"]   = str(data.get("keyword", r.get("keyword", "")))[:500]
            r["location"]  = str(data.get("location", r.get("location", "")))[:300]
            r["min_score"] = int(data.get("min_score", r.get("min_score", 50)))
            r["max_jobs"]  = int(data.get("max_jobs",  r.get("max_jobs",  10)))
            r["max_age_days"] = int(data.get("max_age_days", r.get("max_age_days", 0)) or 0)
            save_user_reminders(uid, reminders)
            return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "Routine not found"}), 404


@app.route("/api/reminders/<reminder_id>", methods=["DELETE"])
def delete_reminder(reminder_id):
    uid = require_user_id()
    if isinstance(uid, tuple):
        return uid
    from database import get_user_reminders, save_user_reminders
    reminders = get_user_reminders(uid)
    updated = [r for r in reminders if str(r.get("id")) != reminder_id]
    if len(updated) == len(reminders):
        return jsonify({"ok": False, "error": "Routine not found"}), 404
    save_user_reminders(uid, updated)
    return jsonify({"ok": True})


@app.route("/api/scheduler/jobs")
def scheduler_jobs():
    """List all scheduled jobs and their next run times."""
    if not _scheduler:
        return jsonify({"ok": False, "error": "Scheduler not running"})
    jobs = []
    for job in _scheduler.get_jobs():
        jobs.append({
            "id": job.id,
            "name": job.name,
            "next_run": str(job.next_run_time) if job.next_run_time else None,
        })
    return jsonify({"ok": True, "jobs": jobs})



# ---------------------------------------------------------------------------
# Setup / Profile (career-ops integration)
# ---------------------------------------------------------------------------

CAREER_OPS_DIR  = os.path.join(BASE_DIR, "tmp_career_ops")
PROFILE_YML     = os.path.join(CAREER_OPS_DIR, "config", "profile.yml")
CV_MD           = os.path.join(CAREER_OPS_DIR, "cv.md")


def _load_profile() -> dict:
    """Load profile.yml as a dict with safe nested defaults."""
    import yaml
    if not os.path.exists(PROFILE_YML):
        return {}
    with open(PROFILE_YML, "r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def _save_profile(data: dict) -> None:
    import yaml
    os.makedirs(os.path.dirname(PROFILE_YML), exist_ok=True)
    with open(PROFILE_YML, "w", encoding="utf-8") as f:
        yaml.dump(data, f, allow_unicode=True, default_flow_style=False, sort_keys=False)


def _nested(d: dict, *keys, default=None):
    """Safe nested dict access."""
    for k in keys:
        if not isinstance(d, dict):
            return default
        d = d.get(k, default)
    return d if d is not None else default


@app.route("/setup", methods=["GET", "POST"])
def setup_profile():
    import yaml

    profile = _load_profile()

    # Ensure nested dicts exist for template rendering
    profile.setdefault("candidate", {})
    profile.setdefault("narrative", {})
    profile.setdefault("target_roles", {})
    profile["target_roles"].setdefault("locations", {})
    profile.setdefault("compensation", {})

    cv_content = ""
    if os.path.exists(CV_MD):
        with open(CV_MD, "r", encoding="utf-8") as f:
            cv_content = f.read()

    # Detect completeness for the status dots
    todo_markers = ["TODO", ""]
    profile_complete = all([
        profile["candidate"].get("full_name", "TODO") not in todo_markers,
        profile["candidate"].get("email", "TODO") not in todo_markers,
        _nested(profile, "narrative", "exit_story", default="TODO") not in todo_markers,
        _nested(profile, "compensation", "target_range", default="TODO") not in todo_markers,
    ])
    cv_complete = bool(cv_content) and "TODO" not in cv_content

    if request.method == "POST":
        tab = request.form.get("tab", "profile")

        if tab == "profile":
            profile["candidate"]["full_name"]    = request.form.get("full_name", "").strip()
            profile["candidate"]["email"]        = request.form.get("email", "").strip()
            profile["candidate"]["phone"]        = request.form.get("phone", "").strip()
            profile["candidate"]["location"]     = request.form.get("location", "").strip()
            profile["candidate"]["linkedin"]     = request.form.get("linkedin", "").strip()
            profile["candidate"]["portfolio_url"]= request.form.get("portfolio_url", "").strip()

            profile["narrative"]["headline"]     = request.form.get("headline", "").strip()
            profile["narrative"]["exit_story"]   = request.form.get("exit_story", "").strip()
            superpowers_raw = request.form.get("superpowers", "")
            profile["narrative"]["superpowers"]  = [s.strip() for s in superpowers_raw.split(",") if s.strip()]

            # Proof points (parallel arrays from form)
            names   = request.form.getlist("proof_name[]")
            urls    = request.form.getlist("proof_url[]")
            metrics = request.form.getlist("proof_metric[]")
            proof_points = []
            for n, u, m in zip(names, urls, metrics):
                if n.strip():
                    proof_points.append({"name": n.strip(), "url": u.strip(), "hero_metric": m.strip()})
            profile["narrative"]["proof_points"] = proof_points

            _save_profile(profile)
            flash("Profile saved successfully.", "success")
            return redirect(url_for("setup_profile") + "?tab=profile")

        elif tab == "targets":
            primary_raw = request.form.get("primary_roles", "")
            profile["target_roles"]["primary"] = [r.strip() for r in primary_raw.splitlines() if r.strip()]

            locs_raw = request.form.get("locations", "")
            profile["target_roles"]["locations"]["preferred"] = [l.strip() for l in locs_raw.splitlines() if l.strip()]

            ind_raw = request.form.get("industries", "")
            profile["target_roles"]["industries"] = [i.strip() for i in ind_raw.splitlines() if i.strip()]

            db_raw = request.form.get("deal_breakers", "")
            profile["deal_breakers"] = [d.strip() for d in db_raw.splitlines() if d.strip()]

            _save_profile(profile)
            flash("Target roles saved.", "success")
            return redirect(url_for("setup_profile") + "?tab=targets")

        elif tab == "comp":
            profile["compensation"]["target_range"]         = request.form.get("target_range", "").strip()
            profile["compensation"]["currency"]             = request.form.get("currency", "INR").strip()
            profile["compensation"]["minimum"]              = request.form.get("minimum", "").strip()
            profile["compensation"]["location_flexibility"] = request.form.get("location_flexibility", "").strip()

            _save_profile(profile)
            flash("Compensation details saved.", "success")
            return redirect(url_for("setup_profile") + "?tab=comp")

        elif tab == "cv":
            cv_text = request.form.get("cv_content", "")
            os.makedirs(os.path.dirname(CV_MD), exist_ok=True)
            with open(CV_MD, "w", encoding="utf-8") as f:
                f.write(cv_text)
            flash("CV saved successfully.", "success")
            return redirect(url_for("setup_profile") + "?tab=cv")

        return redirect(url_for("setup_profile"))

    return render_template(
        "setup.html",
        profile=profile,
        cv_content=cv_content,
        profile_complete=profile_complete,
        cv_complete=cv_complete,
    )


# ---------------------------------------------------------------------------
# Run
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    # Use FLASK_DEBUG env var so LaunchAgent can force production mode.
    # Default to debug=True only when not set (i.e. manual terminal run).
    import os as _os
    _debug = _os.environ.get("FLASK_DEBUG", "1") == "1"
    app.run(debug=_debug, port=int(os.environ.get("PORT", 5002)))
